# -*- coding: utf-8 -*-
"""Gera Auditoria Operacional dos Fluxos Reais do Módulo de Agendamentos (DOCX)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AJUSTES" / "Checklist" / "Auditoria_Operacional_Agendamentos_STP.docx"

# Dados reais extraídos do banco em 29/07/2026
URL = "http://127.0.0.1:5022/transporte/agendamentos"
LOGIN = "http://127.0.0.1:5022/transporte/login"

PACIENTES = [
    (5, "JOSÉ ANTONIO AMBROSIO", "05/07/1953", 73),
    (4, "CICERO DOS SANTOS SILVA", "12/05/1963", 63),
    (3, "HADASSA VITORIA CALIXTO DA LUZ", "06/07/2009", 17),
    (2, "SERVIÇO DE SAÚDE", "01/01/1900", 126),
    (1, "CRISTINA STRAZZACAPPA", "02/09/1965", 60),
    (1993, "SAMUEL DOS SANTOS BARBOSA", "27/05/2015", 11),
    (505, "ADILSON VALLERIO", "—", "—"),
]
ACOMP = [
    (1, "SANDRA REGINA DOS SANTOS BARBOSA", "mãe", "paciente Samuel id=1993"),
]
ESP = [
    "ORTOPEDIA",
    "CARDIOLOGIA",
    "NEUROLOGIA",
    "OFTALMOLOGIA",
    "ONCOLOGIA",
    "CLÍNICA MÉDICA",
    "PEDIATRIA",
]
DEST = [
    "CAMPINAS/HC",
    "CAMPINAS/AME",
    "JAU/HC",
    "BAURU/HC",
    "SOROCABA/CONJUNTO HOSPITALAR",
    "BOTUCATU/HC",
]
MOT = [
    (1, "MARCIO DOS SANTOS CAMARGO"),
    (26, "LENIZA S.T"),
    (27, "CLAUDEMIR S.T"),
]
VEI = [
    (1, "STR-6I20", "ÔNIBUS"),
    (2, "STR-0I87", "ÔNIBUS"),
    (3, "GIL-0A55", "ÔNIBUS"),
    (4, "FZW-4J25", "VAN"),
    (5, "QOE-0A01", "VAN"),
]
FRO = [
    (1, "F00100", "NI Frota Teste Imp"),
    (2, "F00888", "NI Frota Teste 888"),
    (3, "F00020", "NI FROTA 20"),
]
AGS = [
    (16983, "em_andamento", "SAMUEL DOS SANTOS BARBOSA", "LENIZA S.T", "frota F00020", "CAMPINAS/HC", "ORTOPEDIA"),
    (16982, "agendado", "SERVIÇO DE SAÚDE", "MARCIO DOS SANTOS CAMARGO", "STR-6I20", "CAMPINAS/AME", "CARDIOLOGIA"),
    (16981, "agendado", "HADASSA VITORIA…", "—", "—", "JAU/HC", "NEUROLOGIA"),
    (16980, "agendado", "CICERO DOS SANTOS SILVA", "—", "—", "BAURU/HC", "OFTALMOLOGIA"),
    (16979, "agendado", "JOSÉ ANTONIO AMBROSIO", "—", "—", "CAMPINAS/HC", "ORTOPEDIA"),
    (16971, "concluido", "ADILSON VALLERIO", "MARCIO…", "STR-6I20", "CAMPINAS/HC", "ONCOLOGIA"),
]


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_p(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def case_block(doc, num, titulo, objetivo, inicio, passos, esperado, confirmacao, dados_reais, observacoes=None):
    add_h(doc, f"Caso de Teste {num:02d} – {titulo}", 2)
    add_p(doc, "Objetivo", bold=True)
    add_p(doc, objetivo)
    add_p(doc, "Onde iniciar", bold=True)
    add_p(doc, inicio)
    add_p(doc, "Dados reais a utilizar (já cadastrados)", bold=True)
    add_bullets(doc, dados_reais)
    add_p(doc, "Passo a passo", bold=True)
    for i, passo in enumerate(passos, 1):
        doc.add_paragraph(f"{i}. {passo}", style="List Number")
    add_p(doc, "Resultado esperado", bold=True)
    add_p(doc, esperado)
    add_p(doc, "Como confirmar que o processo foi concluído com sucesso", bold=True)
    add_bullets(doc, confirmacao)
    if observacoes:
        add_p(doc, "Observações operacionais", bold=True)
        add_bullets(doc, observacoes)
    add_p(doc, "Resultado da execução:  ☐ Aprovado   ☐ Reprovado   ☐ Bloqueado", bold=True)
    add_p(doc, "Problemas encontrados (se houver): _______________________________________________")
    add_p(doc, "Executor / Data: ___________________________ / ___/___/______")
    doc.add_paragraph()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    t = doc.add_heading("Auditoria Operacional – Fluxos Reais do Módulo de Agendamentos", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_p(
        doc,
        "Sistema de Transporte de Pacientes (STP) — Secretaria Municipal de Saúde",
        bold=True,
        size=12,
    )
    add_p(doc, "Documento gerado com base exclusiva em dados já cadastrados no banco de produção/local do STP.", size=10)
    add_p(doc, "Data de extração dos dados de referência: 29/07/2026.", size=10)
    add_p(doc, "WhatsApp: manter bloqueado (STP_BLOQUEAR_WHATSAPP=1) durante toda a auditoria.", size=10, color=(153, 0, 0))

    add_h(doc, "1. Objetivo", 1)
    add_p(
        doc,
        "Validar, de ponta a ponta, se os fluxos operacionais do módulo de Agendamentos funcionam "
        "de forma simples, rápida, intuitiva e sem erros — simulando o dia a dia real da Secretaria "
        "Municipal de Saúde. Esta auditoria NÃO utiliza dados fictícios.",
    )
    add_bullets(
        doc,
        [
            "Rota principal: " + URL,
            "Login: " + LOGIN + " (usuário real do sistema, ex.: admin / marcio)",
            "Módulo: Transporte → Agendamentos",
            "Escopo: criar, editar, programar, trocar motorista/veículo, alterar destino, cancelar, "
            "reprogramar, imprimir Cartão do Motorista e Folha Espelho, fluxo completo.",
        ],
    )

    add_h(doc, "2. Regras da auditoria", 1)
    add_bullets(
        doc,
        [
            "Usar somente pacientes, acompanhantes, motoristas, veículos, frotas, especialidades e destinos já cadastrados.",
            "Não cadastrar dados fictícios novos (exceto se um fluxo exigir e estiver documentado).",
            "Não enviar WhatsApp real durante a auditoria.",
            "Para cada caso: registrar Aprovado / Reprovado / Bloqueado e anotar evidências.",
            "Critérios contínuos: sem erro na tela; sem perda de informação; sem duplicidade; listagem atualizada; "
            "impressões consistentes; navegação intuitiva; sem necessidade de F5; mensagens claras.",
        ],
    )

    add_h(doc, "3. Inventário de dados reais disponíveis (referência)", 1)

    add_h(doc, "3.1 Pacientes", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["ID", "Nome", "Nascimento", "Idade calc."]):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
    for pid, nome, nasc, idade in PACIENTES:
        row = table.add_row().cells
        row[0].text = str(pid)
        row[1].text = nome
        row[2].text = nasc
        row[3].text = str(idade)

    add_h(doc, "3.2 Acompanhantes", 2)
    add_bullets(doc, [f"ID {a[0]} — {a[1]} ({a[2]}) — {a[3]}" for a in ACOMP])

    add_h(doc, "3.3 Especialidades (cadastro)", 2)
    add_bullets(doc, ESP)

    add_h(doc, "3.4 Destinos frequentes já usados no sistema", 2)
    add_bullets(doc, DEST)

    add_h(doc, "3.5 Motoristas", 2)
    add_bullets(doc, [f"ID {m[0]} — {m[1]}" for m in MOT])

    add_h(doc, "3.6 Veículos", 2)
    add_bullets(doc, [f"ID {v[0]} — {v[1]} ({v[2]})" for v in VEI])

    add_h(doc, "3.7 Frotas", 2)
    add_bullets(doc, [f"ID {f[0]} — {f[1]} — {f[2]}" for f in FRO])

    add_h(doc, "3.8 Agendamentos de referência (já existentes)", 2)
    t2 = doc.add_table(rows=1, cols=7)
    t2.style = "Table Grid"
    hdr2 = t2.rows[0].cells
    for i, h in enumerate(["ID", "Status", "Paciente", "Motorista", "Veíc./Frota", "Destino", "Esp."]):
        hdr2[i].text = h
        set_cell_shading(hdr2[i], "1F4E79")
        for p in hdr2[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
                r.font.size = Pt(9)
    for ag in AGS:
        row = t2.add_row().cells
        for i, val in enumerate(ag):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    add_p(
        doc,
        "Sugestão de uso: 16979–16981 para criar/editar/programar; 16982 para trocas de motorista/veículo; "
        "16983 para impressões (já programado com frota); 16971 apenas consulta (concluído).",
        size=10,
    )

    add_h(doc, "4. Casos de teste operacionais", 1)

    # CT01
    case_block(
        doc,
        1,
        "Criar um Agendamento",
        "Registrar um novo transporte com paciente, especialidade, origem e destino reais, e ver o registro imediatamente na listagem.",
        f"Menu Transporte → Agendamentos ({URL}) → botão Novo / Criar agendamento.",
        [
            "Abrir a tela de novo agendamento.",
            "Selecionar o paciente ID 5 — JOSÉ ANTONIO AMBROSIO (já cadastrado).",
            "Selecionar a especialidade ORTOPEDIA.",
            "Informar origem (ex.: PONTE NOVA / SMS ou o padrão do município já usado no sistema).",
            "Informar destino CAMPINAS/HC.",
            "Informar data/horário de saída e consulta compatíveis com a operação do dia.",
            "Salvar.",
        ],
        "O agendamento aparece imediatamente na listagem, com paciente, especialidade, destino e status corretos, sem mensagem de erro.",
        [
            "Na listagem, localizar o novo ID gerado (coluna ID).",
            "Abrir o registro e conferir paciente = JOSÉ ANTONIO AMBROSIO, especialidade = ORTOPEDIA, destino = CAMPINAS/HC.",
            "Não deve haver flash/alerta de erro vermelho.",
            "Não deve ser necessário pressionar F5 para ver o registro.",
        ],
        [
            "Paciente: ID 5 — JOSÉ ANTONIO AMBROSIO",
            "Especialidade: ORTOPEDIA",
            "Destino: CAMPINAS/HC",
            "Alternativa: paciente ID 4 — CICERO DOS SANTOS SILVA + OFTALMOLOGIA + BAURU/HC",
        ],
        [
            "Anotar o ID gerado: __________ (usar nos casos seguintes se for o fluxo completo).",
            "Se o sistema exigir acompanhante: usar apenas acompanhantes já vinculados ao paciente.",
        ],
    )

    # CT02
    case_block(
        doc,
        2,
        "Editar um Agendamento",
        "Alterar especialidade, horário e observações de um agendamento existente e ver a atualização imediata.",
        "Listagem de Agendamentos → selecionar um agendamento existente (preferir ID 16981 ou o criado no CT01).",
        [
            "Abrir o agendamento ID 16981 (HADASSA VITORIA… / NEUROLOGIA / JAU/HC) ou o ID criado no CT01.",
            "Alterar a especialidade (ex.: de NEUROLOGIA para CLÍNICA MÉDICA — ambas já cadastradas).",
            "Alterar o horário de saída ou consulta em alguns minutos.",
            "Incluir ou ajustar observação operacional (ex.: “Auditoria operacional – CT02”).",
            "Salvar.",
        ],
        "As alterações aparecem imediatamente na listagem e ao reabrir a tela de edição.",
        [
            "Na listagem, conferir especialidade e horários atualizados sem F5.",
            "Reabrir o mesmo ID e confirmar que os valores salvos batem com o que foi editado.",
            "Demais campos (paciente, destino, motorista/veículo se houver) permanecem intactos.",
        ],
        [
            "Agendamento: ID 16981 ou ID criado no CT01",
            "Especialidades disponíveis: ORTOPEDIA, CARDIOLOGIA, NEUROLOGIA, OFTALMOLOGIA, ONCOLOGIA, CLÍNICA MÉDICA, PEDIATRIA",
        ],
    )

    # CT03
    case_block(
        doc,
        3,
        "Programar o Transporte",
        "Vincular motorista e veículo (ou frota) a um agendamento e habilitar as impressões previstas pela regra de negócio.",
        "Listagem → Programar no agendamento ainda sem programação (ex.: ID 16979, 16980 ou 16981).",
        [
            "Abrir Programar do agendamento ID 16979 (JOSÉ ANTONIO AMBROSIO / CAMPINAS/HC / ORTOPEDIA).",
            "Selecionar motorista ID 1 — MARCIO DOS SANTOS CAMARGO.",
            "Selecionar veículo ID 1 — STR-6I20 (ÔNIBUS) OU frota ID 3 — F00020.",
            "Salvar a programação.",
        ],
        "O agendamento fica programado; os botões de Folha Espelho e Cartão do Motorista ficam habilitados conforme a regra "
        "(motorista + veículo OU frota).",
        [
            "Após salvar, permanecer na tela de Programar (comportamento atual) e ver botões de impressão habilitados.",
            "Na listagem, ícones de Folha Espelho e Cartão devem estar habilitados para o ID.",
            "Reabrir Programar e confirmar motorista/veículo (ou frota) gravados.",
        ],
        [
            "Agendamento sem programação: 16979 / 16980 / 16981",
            "Motorista: ID 1 — MARCIO DOS SANTOS CAMARGO",
            "Veículo: ID 1 — STR-6I20  |  Alternativa frota: ID 3 — F00020",
        ],
        [
            "Regra crítica: Folha Espelho deve habilitar também com frota (sem veículo individual), como no ID 16983.",
        ],
    )

    # CT04
    case_block(
        doc,
        4,
        "Trocar o Motorista",
        "Em um agendamento já programado, trocar apenas o motorista sem alterar os demais dados.",
        "Listagem → Programar / Editar do agendamento já programado (ex.: ID 16982).",
        [
            "Abrir o agendamento ID 16982 (já com motorista MARCIO e veículo STR-6I20).",
            "Trocar o motorista para ID 26 — LENIZA S.T (ou ID 27 — CLAUDEMIR S.T).",
            "Não alterar veículo, destino, especialidade nem horários.",
            "Salvar.",
        ],
        "Somente o motorista muda; paciente, destino, veículo, horários e demais campos permanecem iguais.",
        [
            "Reabrir o registro e conferir motorista novo.",
            "Conferir veículo ainda = STR-6I20, destino = CAMPINAS/AME, especialidade = CARDIOLOGIA.",
            "Na listagem, o motorista exibido deve refletir a troca.",
        ],
        [
            "Agendamento: ID 16982",
            "Motorista atual: ID 1 — MARCIO DOS SANTOS CAMARGO",
            "Novo motorista: ID 26 — LENIZA S.T  ou  ID 27 — CLAUDEMIR S.T",
            "Veículo que deve permanecer: STR-6I20",
        ],
    )

    # CT05
    case_block(
        doc,
        5,
        "Trocar o Veículo",
        "Alterar somente o veículo de um agendamento programado e garantir que o novo veículo apareça em telas e documentos.",
        "Listagem → Programar do agendamento programado (ex.: ID 16982 após CT04, ou ID 16983 se usar frota→veículo).",
        [
            "Abrir o agendamento ID 16982.",
            "Alterar o veículo de STR-6I20 para ID 2 — STR-0I87 (ou ID 4 — FZW-4J25).",
            "Não alterar motorista nem demais campos.",
            "Salvar.",
            "Abrir Folha Espelho e/ou Cartão do Motorista e conferir a placa/veículo.",
        ],
        "O novo veículo aparece na listagem, na tela de programação e nas impressões relacionadas.",
        [
            "Tela Programar: veículo = novo.",
            "Listagem: coluna/veículo atualizado.",
            "Impressão: placa/modelo do novo veículo (não o anterior).",
        ],
        [
            "Agendamento: ID 16982 (ou o programado no CT03)",
            "Veículos reais: STR-6I20, STR-0I87, GIL-0A55, FZW-4J25, QOE-0A01",
        ],
    )

    # CT06
    case_block(
        doc,
        6,
        "Alterar o Destino",
        "Trocar o destino de um agendamento existente e ver a atualização na listagem e nas futuras impressões.",
        "Listagem → Editar agendamento (ex.: ID 16980 ou o do fluxo).",
        [
            "Abrir o agendamento ID 16980 (destino atual BAURU/HC) ou outro existente.",
            "Alterar o destino para SOROCABA/CONJUNTO HOSPITALAR (já usado no sistema) ou BOTUCATU/HC.",
            "Salvar.",
            "Se estiver programado, gerar Folha Espelho e conferir o destino impresso.",
        ],
        "O novo destino aparece na listagem e nas impressões geradas após a alteração.",
        [
            "Listagem: destino atualizado sem F5.",
            "Edição: destino salvo correto.",
            "Impressão (se aplicável): destino novo, não o antigo.",
        ],
        [
            "Agendamento: ID 16980 (CICERO DOS SANTOS SILVA / BAURU/HC)",
            "Destinos reais: CAMPINAS/HC, CAMPINAS/AME, JAU/HC, BAURU/HC, SOROCABA/CONJUNTO HOSPITALAR, BOTUCATU/HC",
        ],
    )

    # CT07
    case_block(
        doc,
        7,
        "Cancelar um Agendamento",
        "Cancelar um agendamento e verificar se o status cancelado fica claro conforme a regra de negócio.",
        "Listagem → selecionar agendamento (preferir um criado só para este teste, ou ID sem impacto operacional crítico).",
        [
            "Selecionar um agendamento elegível (ex.: o criado no CT01, se ainda existir, ou 16979 se não for necessário manter).",
            "Acionar a ação de Cancelar (botão/fluxo previsto na tela).",
            "Confirmar o cancelamento se o sistema pedir confirmação.",
            "Observar status e feedback na tela.",
        ],
        "O sistema identifica claramente que o agendamento foi cancelado (status, cor, rótulo ou mensagem), conforme a regra definida.",
        [
            "Na listagem, o registro exibe status cancelado de forma inequívoca.",
            "Não deve ser possível programar/imprimir como se estivesse ativo (conforme regra).",
            "Não deve haver erro ou sumiço silencioso do registro (a menos que a regra oculte cancelados — anotar o comportamento).",
        ],
        [
            "Usar preferencialmente agendamento de teste (criado no CT01/CT13).",
            "Evitar cancelar ID 16971 (já concluído) e ID 16983 se ainda for necessário para impressões.",
        ],
        [
            "Anotar se cancelados continuam visíveis na listagem padrão ou só com filtro de status.",
        ],
    )

    # CT08
    case_block(
        doc,
        8,
        "Reprogramar um Transporte",
        "Substituir motorista, veículo e horário de um agendamento já programado sem deixar dados inconsistentes.",
        "Listagem → Programar de um agendamento já programado (ex.: ID 16982 ou 16983).",
        [
            "Abrir Programar do ID 16982 (ou 16983).",
            "Alterar motorista (ex.: MARCIO → CLAUDEMIR S.T ID 27).",
            "Alterar veículo (ex.: STR-6I20 → GIL-0A55 ID 3) ou frota.",
            "Se a tela permitir, ajustar horário de saída.",
            "Salvar.",
        ],
        "A programação anterior é substituída corretamente; não restam motorista/veículo antigos misturados; impressões passam a refletir a nova programação.",
        [
            "Reabrir Programar: apenas os novos valores.",
            "Listagem: motorista/veículo novos.",
            "Impressões: dados novos; sem duplicar viagem.",
        ],
        [
            "Agendamento: ID 16982 ou 16983",
            "Motoristas: MARCIO (1), LENIZA (26), CLAUDEMIR (27)",
            "Veículos/frotas: STR-6I20, GIL-0A55, F00020",
        ],
    )

    # CT09
    case_block(
        doc,
        9,
        "Imprimir Cartão do Motorista",
        "Gerar o Cartão do Motorista para um motorista com viagens programadas e validar conteúdo operacional.",
        "Listagem de Agendamentos (filtros por data/motorista) OU rota de impressão do Cartão a partir do agendamento programado.",
        [
            "Selecionar motorista ID 26 — LENIZA S.T (possui viagem no ID 16983) OU motorista ID 1 — MARCIO (ID 16982 / 16971).",
            "Gerar o Cartão do Motorista (ícone/botão na listagem ou na Programar).",
            "Abrir a visualização/PDF/HTML gerado.",
            "Conferir paciente, acompanhante (se houver), horários, destinos, veículo/frota, motorista e ordem das viagens.",
        ],
        "O cartão apresenta dados corretos e coerentes com o(s) agendamento(s) programado(s) do motorista no dia.",
        [
            "Paciente correto (ex.: SAMUEL DOS SANTOS BARBOSA no 16983).",
            "Acompanhante: SANDRA REGINA DOS SANTOS BARBOSA (se vinculado).",
            "Destino/horário/veículo ou frota batem com a tela.",
            "Ordem das viagens cronológica quando houver mais de uma.",
            "Sem erro HTTP / página em branco.",
        ],
        [
            "Referência forte: Agendamento ID 16983 — SAMUEL / LENIZA / frota F00020 / CAMPINAS/HC / ORTOPEDIA",
            "Alternativa: ID 16982 — SERVIÇO DE SAÚDE / MARCIO / STR-6I20 / CAMPINAS/AME",
        ],
    )

    # CT10
    case_block(
        doc,
        10,
        "Motorista com Várias Viagens",
        "Validar cartão/listagem quando o mesmo motorista tem mais de um transporte no mesmo dia.",
        "Filtrar agendamentos do dia pelo motorista escolhido; gerar Cartão do Motorista.",
        [
            "Identificar um motorista com 2+ viagens no mesmo dia (filtrar listagem por data + motorista).",
            "Se no dia da auditoria não houver naturalmente 2+ viagens: programar dois agendamentos existentes "
            "(ex.: 16979 e 16980) para o mesmo motorista (MARCIO ID 1) e mesma data, horários diferentes.",
            "Gerar o Cartão do Motorista desse motorista/data.",
            "Conferir se todas as viagens aparecem, em ordem cronológica, sem sumiço e sem duplicidade.",
        ],
        "Todas as viagens do dia aparecem uma única vez, em sequência cronológica correta, com impressão coerente.",
        [
            "Contar viagens na listagem filtrada = contar no cartão.",
            "Horários em ordem crescente.",
            "Nenhuma viagem duplicada.",
            "Nenhuma viagem programada ausente.",
        ],
        [
            "Motorista sugerido: ID 1 — MARCIO DOS SANTOS CAMARGO",
            "Agendamentos a programar no mesmo dia (se necessário): 16979 e 16980 (ou 16981)",
            "Veículos distintos ou o mesmo — anotar o comportamento",
        ],
        [
            "Se o motorista não tiver múltiplas viagens no dia, criar a situação com dados reais já existentes (programar 2 IDs).",
            "Não inventar pacientes novos.",
        ],
    )

    # CT11
    case_block(
        doc,
        11,
        "Imprimir Folha Espelho",
        "Gerar a Folha Espelho de um agendamento completo e validar campos e layout frente ao modelo oficial.",
        "Listagem → ícone Folha Espelho do agendamento programado (ex.: ID 16983) OU tela Programar → Folha Espelho.",
        [
            "Selecionar o agendamento ID 16983 (SAMUEL / LENIZA / frota F00020 / CAMPINAS/HC / ORTOPEDIA) — já completo.",
            "Gerar a Folha Espelho.",
            "Conferir paciente, acompanhante, idade, motorista, veículo/frota, origem, destino, especialidade, horários.",
            "Comparar layout com o modelo oficial da Secretaria (folha física/arquivo modelo).",
            "Conferir se o atendente impresso é o usuário logado (não um nome fixo antigo).",
        ],
        "Folha Espelho com todos os campos corretos e layout alinhado ao modelo oficial; elegível com motorista + frota.",
        [
            "Paciente = SAMUEL DOS SANTOS BARBOSA",
            "Idade coerente com data de nascimento 27/05/2015",
            "Acompanhante = SANDRA REGINA… (se constar no vínculo)",
            "Motorista = LENIZA S.T",
            "Frota/veículo = F00020",
            "Destino = CAMPINAS/HC | Especialidade = ORTOPEDIA",
            "Atendente = usuário da sessão",
            "Sem erro; botão habilitado na listagem",
        ],
        [
            "Agendamento principal: ID 16983",
            "Alternativa após programar: ID 16982",
        ],
        [
            "Anexar ou guardar captura/PDF como evidência da auditoria.",
        ],
    )

    # CT12
    case_block(
        doc,
        12,
        "Alterar um Agendamento Já Programado",
        "Modificar apenas um dado de um agendamento que já possui motorista e veículo, garantindo que o restante permanece intacto.",
        "Listagem → Editar / Programar de agendamento programado (ex.: ID 16982 ou 16983).",
        [
            "Abrir o agendamento ID 16982 (ou 16983).",
            "Anotar mentalmente/print: paciente, motorista, veículo/frota, destino, especialidade, horários, status.",
            "Alterar SOMENTE um campo (ex.: observação OU horário OU especialidade — escolher um).",
            "Salvar.",
            "Reabrir e comparar todos os demais campos com o snapshot anterior.",
        ],
        "Apenas o campo alterado muda; motorista, veículo, paciente, destino e demais dados permanecem inalterados.",
        [
            "Diff mental/checklist: 1 campo diferente, N-1 iguais.",
            "Impressões posteriores refletem só a mudança feita (se o campo impresso for o alterado).",
        ],
        [
            "Agendamento: ID 16982 ou 16983",
            "Campos de controle: paciente, motorista, veículo/frota, destino, especialidade, horários, status",
        ],
    )

    # CT13
    case_block(
        doc,
        13,
        "Fluxo Completo (ciclo operacional)",
        "Executar o ciclo completo do transporte — do agendamento ao cancelamento — confirmando consistência em cada etapa.",
        f"Início em {URL} com usuário real logado.",
        [
            "Criar um agendamento com paciente real (ex.: ID 3 — HADASSA VITORIA… ou ID 4 — CICERO…), especialidade e destino reais. Anotar o novo ID.",
            "Editar o agendamento (especialidade ou horário/observação).",
            "Programar motorista (ex.: MARCIO ID 1).",
            "Programar/confirmar veículo (ex.: STR-6I20) ou frota (F00020).",
            "Alterar motorista (ex.: para LENIZA ID 26).",
            "Alterar veículo (ex.: para STR-0I87).",
            "Alterar destino (ex.: para CAMPINAS/AME).",
            "Imprimir Cartão do Motorista e validar dados da etapa atual.",
            "Imprimir Folha Espelho e validar dados da etapa atual.",
            "Cancelar o agendamento e confirmar status cancelado.",
        ],
        "Em cada etapa, a listagem e as telas refletem corretamente o estado; impressões batem com os dados vigentes; ao final o cancelamento está claro e sem resíduos inconsistentes.",
        [
            "Checklist por etapa: após cada passo, conferir listagem + reabertura do registro.",
            "Após impressões: dados = estado pós-alterações (não o estado inicial).",
            "Após cancelar: status inequívoco; impressões/ações respeitam a regra para cancelados.",
            "Nenhum erro de tela em nenhuma etapa.",
        ],
        [
            "Pacientes sugeridos: HADASSA (3), CICERO (4), JOSÉ ANTONIO (5), CRISTINA (1)",
            "Motoristas: MARCIO (1), LENIZA (26), CLAUDEMIR (27)",
            "Veículos/frotas: STR-6I20, STR-0I87, F00020",
            "Destinos: CAMPINAS/HC, CAMPINAS/AME, JAU/HC, BAURU/HC",
        ],
        [
            "Este caso concentra a prova de vida operacional da Secretaria. Priorizar evidências (prints) em cada etapa.",
            "Manter WhatsApp bloqueado.",
        ],
    )

    add_h(doc, "5. Critérios gerais de aprovação (checklist contínuo)", 1)
    crit = [
        "Nenhum erro na tela (mensagem vermelha, 500, traceback).",
        "Nenhuma informação perdida após salvar.",
        "Nenhuma duplicidade de registros.",
        "Atualização imediata da listagem (sem F5).",
        "Regras de negócio respeitadas (programação, impressão, cancelamento, status).",
        "Impressões consistentes com a tela.",
        "Navegação simples e intuitiva para o atendente.",
        "Funcionamento sem necessidade de recarregar a página.",
        "Mensagens claras para o usuário.",
        "IDs visíveis e úteis para rastreio operacional.",
        "WhatsApp não disparado durante a auditoria.",
    ]
    for c in crit:
        doc.add_paragraph(f"☐ {c}")

    add_h(doc, "6. Relatório final da auditoria", 1)
    add_p(
        doc,
        "Preencher ao término da execução dos 13 casos. Usar dados e evidências reais observados.",
    )

    add_h(doc, "6.1 Resumo executivo", 2)
    add_p(doc, "Data da auditoria: ___/___/______")
    add_p(doc, "Executor(es): ________________________________")
    add_p(doc, "Ambiente: http://127.0.0.1:5022/transporte (local) / ☐ outro: ________")
    add_p(doc, "Usuário logado: ____________________")
    add_p(doc, "Resultado geral:  ☐ Aprovado  ☐ Aprovado com ressalvas  ☐ Reprovado")

    add_h(doc, "6.2 Fluxos aprovados", 2)
    add_p(doc, "Listar os casos APROVADOS (números e títulos):")
    add_p(doc, "_________________________________________________________________")
    add_p(doc, "_________________________________________________________________")

    add_h(doc, "6.3 Fluxos reprovados", 2)
    add_p(doc, "Listar os casos REPROVADOS (números e títulos):")
    add_p(doc, "_________________________________________________________________")
    add_p(doc, "_________________________________________________________________")

    add_h(doc, "6.4 Problemas encontrados", 2)
    add_p(
        doc,
        "Para cada problema, preencher o quadro abaixo (copiar o bloco quantas vezes for necessário).",
    )

    for i in range(1, 6):
        add_p(doc, f"Problema #{i}", bold=True)
        add_bullets(
            doc,
            [
                "Descrição: _______________________________________________________________",
                "Localização exata (tela/rota/ID do agendamento/campo): ____________________",
                "Caso de teste relacionado (01–13): ____",
                "Passos para reproduzir: __________________________________________________",
                "Impacto operacional (atendente/paciente/motorista/documento): ______________",
                "Correção sugerida: _______________________________________________________",
                "Prioridade:  ☐ Crítica  ☐ Alta  ☐ Média  ☐ Baixa",
                "Evidência (print/PDF/ID): ________________________________________________",
            ],
        )

    add_h(doc, "6.5 Matriz rápida CT × Resultado", 2)
    tm = doc.add_table(rows=1, cols=4)
    tm.style = "Table Grid"
    for i, h in enumerate(["CT", "Título", "Resultado", "Obs."]):
        tm.rows[0].cells[i].text = h
        set_cell_shading(tm.rows[0].cells[i], "1F4E79")
        for p in tm.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
    titulos = [
        "Criar agendamento",
        "Editar agendamento",
        "Programar transporte",
        "Trocar motorista",
        "Trocar veículo",
        "Alterar destino",
        "Cancelar agendamento",
        "Reprogramar transporte",
        "Imprimir Cartão Motorista",
        "Motorista várias viagens",
        "Imprimir Folha Espelho",
        "Alterar já programado",
        "Fluxo completo",
    ]
    for i, tit in enumerate(titulos, 1):
        row = tm.add_row().cells
        row[0].text = f"{i:02d}"
        row[1].text = tit
        row[2].text = "☐ Apr ☐ Rep ☐ Bloq"
        row[3].text = ""

    add_h(doc, "6.6 Conclusão e recomendação", 2)
    add_p(doc, "_________________________________________________________________")
    add_p(doc, "_________________________________________________________________")
    add_p(doc, "_________________________________________________________________")
    add_p(doc, "Assinatura do responsável pela auditoria: ___________________________")

    add_h(doc, "7. Observações finais", 1)
    add_bullets(
        doc,
        [
            "Todos os IDs e nomes deste documento foram extraídos do banco local do STP na data de geração.",
            "Se algum registro for alterado/cancelado durante a auditoria, anotar o novo estado — não inventar substitutos fictícios.",
            "Após a auditoria operacional e correções, avaliar remoção do bloqueio de WhatsApp somente com aprovação explícita.",
            "Documento de apoio paralelo: Checklist_Auditoria_Agendamentos_STP.docx (checklist detalhado de UI/regras).",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Gerado: {OUT}")


if __name__ == "__main__":
    main()
