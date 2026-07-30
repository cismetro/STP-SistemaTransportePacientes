# -*- coding: utf-8 -*-
"""Gera/atualiza AJUSTES/Checklist/Checklist_Validacao_STP.docx

Usa DADOS REAIS do banco local (db/transporte_pacientes.db).
Regra: para acompanhante, BUSCAR paciente já cadastrado — não inventar paciente novo.
"""
from __future__ import annotations

import sqlite3
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DB_PATH = ROOT / "db" / "transporte_pacientes.db"
OUT = ROOT / "AJUSTES" / "Checklist" / "Checklist_Validacao_STP.docx"

_BOX_ID = 2000


def _next_box_id() -> int:
    global _BOX_ID
    _BOX_ID += 1
    return _BOX_ID


def add_clickable_checkbox(paragraph, name: str, box_id: int | None = None) -> None:
    if box_id is None:
        box_id = _next_box_id()
    safe_name = (
        name.replace("&", "&amp;")
        .replace('"', "&quot;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    xml = (
        '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
        "<w:sdtPr>"
        f'<w:alias w:val="{safe_name}"/>'
        f'<w:tag w:val="{safe_name}"/>'
        f'<w:id w:val="{box_id}"/>'
        "<w14:checkbox>"
        '<w14:checked w14:val="0"/>'
        '<w14:checkedState w14:val="2612" w14:font="Segoe UI Symbol"/>'
        '<w14:uncheckedState w14:val="2610" w14:font="Segoe UI Symbol"/>'
        "</w14:checkbox>"
        "</w:sdtPr>"
        "<w:sdtContent>"
        "<w:r><w:rPr>"
        '<w:rFonts w:ascii="Segoe UI Symbol" w:hAnsi="Segoe UI Symbol"/>'
        '<w:sz w:val="22"/>'
        f"</w:rPr><w:t>☐</w:t></w:r>"
        "</w:sdtContent>"
        "</w:sdt>"
    )
    paragraph._p.append(parse_xml(xml))


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shading)


def set_run_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, *, size=11, bold=False, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_narrow_margins(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)


def style_heading(paragraph, level=1):
    paragraph.style = f"Heading {level}"
    for run in paragraph.runs:
        run.font.name = "Calibri"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)


def add_checklist_table(doc, items, start_num=1):
    headers = ["#", "Funcionalidade principal", "OK", "Não OK", "Revisar", "Observações"]
    widths = [Cm(0.9), Cm(9.2), Cm(1.3), Cm(1.6), Cm(1.6), Cm(3.0)]
    table = doc.add_table(rows=1 + len(items), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color=(255, 255, 255))
        set_cell_shading(hdr[i], "1A3A5C")
        hdr[i].width = widths[i]

    checkbox_cols = {2: "ok", 3: "nao_ok", 4: "revisar"}

    for idx, text in enumerate(items):
        item_num = start_num + idx
        row = table.rows[idx + 1].cells
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
        ]
        for c in range(6):
            row[c].text = ""
            p = row[c].paragraphs[0]
            p.alignment = aligns[c]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if c == 0:
                run = p.add_run(str(item_num))
                set_run_font(run, size=10, bold=True)
            elif c == 1:
                run = p.add_run(text)
                set_run_font(run, size=10, bold=False)
            elif c in checkbox_cols:
                add_clickable_checkbox(p, f"stp_item{item_num}_{checkbox_cols[c]}")
            else:
                run = p.add_run("")
                set_run_font(run, size=10)
            row[c].width = widths[c]
        if idx % 2 == 1:
            for c in range(6):
                set_cell_shading(row[c], "F0F4F8")

    return start_num + len(items)


def add_header_fields(doc):
    add_para(doc, "Dados do teste", size=12, bold=True, space_before=6, space_after=4)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Responsável:", "_______________________________________________"),
        ("Data do teste:", "____/____/________"),
    ]
    for i, (lab, val) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        p0 = c0.paragraphs[0]
        set_run_font(p0.add_run(lab), size=10, bold=True)
        c0.width = Cm(3.5)
        set_cell_shading(c0, "E8EEF4")
        c1.text = ""
        p1 = c1.paragraphs[0]
        set_run_font(p1.add_run(val), size=10)
        c1.width = Cm(14.1)

    c0, c1 = table.rows[2].cells
    c0.text = ""
    p0 = c0.paragraphs[0]
    set_run_font(p0.add_run("Ambiente:"), size=10, bold=True)
    c0.width = Cm(3.5)
    set_cell_shading(c0, "E8EEF4")
    c1.text = ""
    p1 = c1.paragraphs[0]
    add_clickable_checkbox(p1, "stp_ambiente_homologacao")
    set_run_font(p1.add_run(" Homologação     "), size=10)
    add_clickable_checkbox(p1, "stp_ambiente_producao")
    set_run_font(p1.add_run(" Produção"), size=10)
    c1.width = Cm(14.1)
    doc.add_paragraph()


def add_conclusao(doc):
    h = doc.add_heading("Conclusão / Parecer", level=2)
    style_heading(h, 2)

    add_para(doc, "Resultado geral:", size=11, bold=True, space_after=4)
    resultado = doc.add_paragraph()
    resultado.paragraph_format.space_after = Pt(8)
    add_clickable_checkbox(resultado, "stp_aprovado")
    set_run_font(resultado.add_run(" Aprovado           "), size=10)
    add_clickable_checkbox(resultado, "stp_aprovado_ressalvas")
    set_run_font(resultado.add_run(" Aprovado com ressalvas           "), size=10)
    add_clickable_checkbox(resultado, "stp_reprovado")
    set_run_font(resultado.add_run(" Reprovado"), size=10)

    add_para(doc, "Quantitativo dos itens avaliados:", size=11, bold=True, space_after=4)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    for i, lab in enumerate(
        [
            "Quantidade de itens OK",
            "Quantidade de itens Não OK",
            "Quantidade de itens Revisar",
            "Total de itens avaliados",
        ]
    ):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        p0 = c0.paragraphs[0]
        set_run_font(p0.add_run(lab), size=10, bold=True)
        set_cell_shading(c0, "E8EEF4")
        c0.width = Cm(7)
        c1.text = ""
        p1 = c1.paragraphs[0]
        set_run_font(p1.add_run("____________________"), size=10)
        c1.width = Cm(10.6)

    doc.add_paragraph()
    add_para(doc, "Observações finais / ressalvas:", size=11, bold=True, space_after=4)
    for _ in range(3):
        add_para(doc, "_" * 95, size=10, space_after=2)

    add_para(doc, "Assinaturas", size=11, bold=True, space_before=10, space_after=8)
    sig = doc.add_table(rows=3, cols=2)
    sig.style = "Table Grid"
    labels = [
        ("Testador", "Responsável técnico"),
        ("Nome: _______________________________", "Nome: _______________________________"),
        ("Assinatura: _________________________", "Assinatura: _________________________"),
    ]
    for i, (a, b) in enumerate(labels):
        ca, cb = sig.rows[i].cells
        for cell, text in ((ca, a), (cb, b)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, size=10, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "1A3A5C")
                run.font.color.rgb = RGBColor(255, 255, 255)

    add_para(doc, "Data: ____/____/________", size=10, space_before=8, space_after=4)


def _fmt_data(val):
    if not val:
        return "—"
    s = str(val)[:10]
    if len(s) == 10 and s[4] == "-":
        y, m, d = s.split("-")
        return f"{d}/{m}/{y}"
    return s


def carregar_dados_reais():
    """Lê cadastros reais do banco STP para montar o roteiro de teste."""
    if not DB_PATH.exists():
        raise FileNotFoundError(f"Banco não encontrado: {DB_PATH}")

    con = sqlite3.connect(str(DB_PATH))
    con.row_factory = sqlite3.Row
    cur = con.cursor()

    def one(sql, params=()):
        return cur.execute(sql, params).fetchone()

    def many(sql, params=()):
        return cur.execute(sql, params).fetchall()

    # Preferir pacientes conhecidos / com condição; senão primeiros ativos
    preferidos = [
        "SAMUEL DOS SANTOS BARBOSA",
        "MARIA APARECIDA FERREIRA",
        "ANTONIO CARLOS MENDES",
        "ADEMAR DE SOUZA AMARAL",
        "ADAIL FRANCISCO SOARES",
    ]
    pacientes = []
    for nome in preferidos:
        row = one(
            "SELECT id, nome, cpf, telefone, data_nascimento, endereco, observacoes, "
            "condicao_paciente FROM pacientes WHERE ativo=1 AND UPPER(nome)=UPPER(?) LIMIT 1",
            (nome,),
        )
        if row:
            pacientes.append(row)

    if len(pacientes) < 4:
        extras = many(
            "SELECT id, nome, cpf, telefone, data_nascimento, endereco, observacoes, "
            "condicao_paciente FROM pacientes WHERE ativo=1 AND nome IS NOT NULL "
            "AND TRIM(nome)!='' AND cpf IS NOT NULL ORDER BY nome LIMIT 8"
        )
        ids = {p["id"] for p in pacientes}
        for e in extras:
            if e["id"] not in ids:
                pacientes.append(e)
            if len(pacientes) >= 4:
                break

    motorista = one(
        "SELECT id, nome, cpf, cnh, categoria_cnh, vencimento_cnh, status "
        "FROM motoristas WHERE status='ativo' AND cpf NOT LIKE '900.%' "
        "ORDER BY nome LIMIT 1"
    )
    veiculo = one(
        "SELECT id, placa, marca, modelo, tipo, capacidade, numero_frota, frota_id "
        "FROM veiculos WHERE ativo=1 AND capacidade >= 6 "
        "ORDER BY CASE WHEN capacidade=6 THEN 0 ELSE 1 END, capacidade, placa LIMIT 1"
    )
    if veiculo is None:
        veiculo = one(
            "SELECT id, placa, marca, modelo, tipo, capacidade, numero_frota, frota_id "
            "FROM veiculos WHERE ativo=1 ORDER BY placa LIMIT 1"
        )
    frota = one("SELECT id, numero, nome FROM frotas WHERE ativo=1 ORDER BY id LIMIT 1")

    total_pac = one("SELECT COUNT(*) AS n FROM pacientes WHERE ativo=1")["n"]
    total_acomp = one(
        "SELECT COUNT(*) AS n FROM acompanhantes WHERE ativo=1 OR ativo IS NULL"
    )["n"]

    # Acompanhante de teste (novo) — sempre vinculado a paciente JÁ existente
    acomp_nome = "SANDRA DOS SANTOS BARBOSA"
    acomp_parentesco = "Mãe"
    acomp_tel = "(19) 98320-1356"

    # Se já existir acompanhante no banco para o paciente A, usar o real
    pac_a = pacientes[0] if pacientes else None
    acomp_existente = None
    if pac_a:
        acomp_existente = one(
            "SELECT id, nome, parentesco, telefone FROM acompanhantes "
            "WHERE paciente_id=? AND (ativo=1 OR ativo IS NULL) ORDER BY id LIMIT 1",
            (pac_a["id"],),
        )
    if acomp_existente:
        acomp_nome = acomp_existente["nome"]
        acomp_parentesco = acomp_existente["parentesco"] or acomp_parentesco
        acomp_tel = acomp_existente["telefone"] or acomp_tel

    con.close()

    return {
        "total_pacientes": total_pac,
        "total_acompanhantes": total_acomp,
        "pacientes": pacientes,
        "motorista": motorista,
        "veiculo": veiculo,
        "frota": frota,
        "acompanhante_novo": {
            "nome": acomp_nome,
            "parentesco": acomp_parentesco,
            "telefone": acomp_tel,
            "ja_existe": bool(acomp_existente),
        },
    }


def add_exemplo_cadastro(doc, dados):
    add_para(
        doc,
        "Dados REAIS do sistema (buscar no STP — não inventar paciente novo)",
        size=12,
        bold=True,
        space_before=4,
        space_after=4,
    )
    add_para(
        doc,
        f"Fonte: banco local transporte_pacientes.db | Pacientes ativos: {dados['total_pacientes']} | "
        f"Acompanhantes cadastrados: {dados['total_acompanhantes']}. "
        "REGRA: o acompanhante só pode ser cadastrado vinculado a um paciente JÁ existente. "
        "No cadastro de acompanhante, pesquise/selecione o paciente pelo nome ou CPF abaixo.",
        size=9,
        space_after=6,
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(["Cadastro / papel", "Como buscar / dados reais no sistema"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_run_font(p.add_run(h), size=9, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, "1A3A5C")
        cell.width = Cm(4.2) if i == 0 else Cm(13.4)

    pacs = dados["pacientes"]
    mot = dados["motorista"]
    vei = dados["veiculo"]
    fro = dados["frota"]
    ac = dados["acompanhante_novo"]

    def pac_line(p, papel):
        cond = p["condicao_paciente"] or "—"
        obs = (p["observacoes"] or "—")
        if len(str(obs)) > 80:
            obs = str(obs)[:77] + "..."
        return (
            f"{papel}: BUSCAR paciente id {p['id']} — {p['nome']} | CPF {p['cpf']} | "
            f"tel {p['telefone'] or '—'} | nasc {_fmt_data(p['data_nascimento'])} | "
            f"condição {cond} | obs {obs}"
        )

    linhas = []
    if fro:
        linhas.append(
            (
                "Frota (já cadastrada)",
                f"BUSCAR frota id {fro['id']} — {fro['numero']} / {fro['nome']}",
            )
        )
    if vei:
        linhas.append(
            (
                "Veículo (já cadastrado)",
                f"BUSCAR placa {vei['placa']} — {vei['marca']} {vei['modelo']} | "
                f"tipo {vei['tipo']} | capacidade {vei['capacidade']} | "
                f"nº frota Access {vei['numero_frota'] or '—'}",
            )
        )
    if mot:
        linhas.append(
            (
                "Motorista (já cadastrado)",
                f"BUSCAR id {mot['id']} — {mot['nome']} | CPF {mot['cpf']} | "
                f"CNH {mot['cnh']} cat {mot['categoria_cnh']} | "
                f"validade {_fmt_data(mot['vencimento_cnh'])} | {mot['status']}",
            )
        )

    papeis = ["Paciente A (COM acompanhante)", "Paciente B (SEM acompanhante)",
              "Paciente C", "Paciente D"]
    for i, papel in enumerate(papeis):
        if i < len(pacs):
            linhas.append((papel, pac_line(pacs[i], "Já existe")))

    if pacs:
        acao = (
            "JÁ EXISTE no banco — usar este registro"
            if ac["ja_existe"]
            else "CADASTRAR AGORA vinculando ao Paciente A (buscar paciente pelo CPF/nome)"
        )
        linhas.append(
            (
                "Acompanhante do Paciente A",
                f"{acao}: {ac['nome']} | parentesco {ac['parentesco']} | "
                f"tel {ac['telefone']} | paciente: {pacs[0]['nome']} (CPF {pacs[0]['cpf']})",
            )
        )

    linhas.append(
        (
            "Agendamento (base)",
            "Data de teste: 30/07/2026 | Saída 07:30 | Consulta 09:00 | "
            "Destino: Hospital Estadual — Campinas/SP | "
            f"Motorista: {mot['nome'] if mot else '—'} | "
            f"Veículo: {vei['placa'] if vei else '—'}",
        )
    )

    for idx, (cadastro, texto) in enumerate(linhas):
        row = table.add_row().cells
        row[0].text = ""
        set_run_font(row[0].paragraphs[0].add_run(cadastro), size=9, bold=True)
        row[0].width = Cm(4.2)
        row[1].text = ""
        set_run_font(row[1].paragraphs[0].add_run(texto), size=9)
        row[1].width = Cm(13.4)
        if idx % 2 == 1:
            set_cell_shading(row[0], "F0F4F8")
            set_cell_shading(row[1], "F0F4F8")

    doc.add_paragraph()


def montar_secoes(dados):
    pacs = dados["pacientes"]
    mot = dados["motorista"]
    vei = dados["veiculo"]
    fro = dados["frota"]
    ac = dados["acompanhante_novo"]

    p_a = pacs[0] if len(pacs) > 0 else None
    p_b = pacs[1] if len(pacs) > 1 else None
    p_c = pacs[2] if len(pacs) > 2 else None
    p_d = pacs[3] if len(pacs) > 3 else None

    n_a = p_a["nome"] if p_a else "Paciente A"
    cpf_a = p_a["cpf"] if p_a else "—"
    id_a = p_a["id"] if p_a else "—"
    n_b = p_b["nome"] if p_b else "Paciente B"
    n_c = p_c["nome"] if p_c else "Paciente C"
    n_d = p_d["nome"] if p_d else "Paciente D"
    n_mot = mot["nome"] if mot else "Motorista"
    placa = vei["placa"] if vei else "—"
    cap = vei["capacidade"] if vei else 6
    frota_txt = f"{fro['numero']} / {fro['nome']}" if fro else "frota cadastrada"
    ac_nome = ac["nome"]

    if ac["ja_existe"]:
        item_acomp = (
            f"Acompanhante: localizar {ac_nome} já vinculado a {n_a} (CPF {cpf_a}) "
            f"e conferir parentesco/telefone."
        )
    else:
        item_acomp = (
            f"Acompanhante: em Cadastros → Acompanhantes, BUSCAR o paciente {n_a} "
            f"(id {id_a} / CPF {cpf_a}) e cadastrar {ac_nome} "
            f"(parentesco {ac['parentesco']}, tel {ac['telefone']}) VINCULADO a esse paciente. "
            "Não criar paciente novo."
        )

    return [
        (
            "1. Cadastros",
            [
                f"Paciente A: BUSCAR no sistema {n_a} (id {id_a}, CPF {cpf_a}) — "
                "conferir telefone, condição/observações. Não cadastrar outro paciente com o mesmo CPF.",
                item_acomp,
                f"Editar paciente {n_a}: se já houver acompanhante ativo, a ficha NÃO pede novo AC "
                "ao salvar o paciente; formulário de novo AC fica recolhido (botão "
                "'Adicionar outro acompanhante').",
                f"Editar paciente {n_a}: na tabela de acompanhantes deve aparecer só 1 registro ativo "
                f"de {ac_nome} (sem duplicata incompleta).",
                f"Motorista: BUSCAR {n_mot} — conferir CNH, categoria, validade e situação Ativo.",
                f"Frota: localizar {frota_txt} e conferir o ícone (ℹ️) informativo na aba Frota.",
                "Frota: orientação explica frota como conjunto operacional e veículo vinculado.",
                f"Veículo: BUSCAR placa {placa} — conferir modelo, tipo, capacidade {cap} e situação.",
                "Ícone ℹ️ presente e padronizado em Paciente, Acompanhante, Motorista e Veículos/Frota.",
            ],
        ),
        (
            "1b. Idade (Paciente e Acompanhante)",
            [
                "Cadastro/edição de Paciente: existe somente Data de Nascimento (dia/mês/ano); "
                "NÃO há campo Idade digitável nem preview de idade no formulário.",
                "Cadastro/edição de Acompanhante: existe somente Data de Nascimento; "
                "NÃO há campo Idade no formulário.",
                "Listagem de Pacientes: coluna Idade calculada automaticamente a partir da data "
                "de nascimento (anos completos; menores de 1 ano em meses/dias).",
                f"Listagem de Acompanhantes: colunas Idade (AC) e Idade Pac. — conferir {ac_nome} "
                f"vinculado a {n_a}.",
                "Impressão de pacientes: inclui Idade calculada (não armazenada no banco).",
                "Cartão do Motorista e Folha Espelho: IDADE do paciente e do AC calculadas "
                "na data do agendamento (mesma rotina central validadores/idade.py).",
                "Negativo: tentar data de nascimento futura ou anterior a 01/01/1850 → sistema rejeita.",
            ],
        ),
        (
            "2. Agendamento",
            [
                f"Criar viagem 30/07/2026 (saída 07:30 / consulta 09:00) com paciente {n_a}, "
                f"acompanhante {ac_nome}, destino Hospital Estadual (Campinas/SP).",
                f"Programar com motorista {n_mot} e veículo {placa} (frota {frota_txt}).",
            ],
        ),
        (
            "3. Capacidade do veículo",
            [
                f"Veículo {placa} capacidade {cap}: lotação = pacientes + acompanhantes não pode passar de {cap}.",
                f"Exemplo permitido: se capacidade={cap}, montar ocupação <= {cap} → sistema permite.",
                f"Exemplo bloqueado: montar ocupação > {cap} → sistema bloqueia com mensagem clara.",
            ],
        ),
        (
            "4. Disponibilidade do motorista",
            [
                f"Com {n_mot} já programado às 07:30 de 30/07/2026, tentar outro agendamento "
                "no mesmo horário → conflito detectado e bloqueado.",
            ],
        ),
        (
            "5. Disponibilidade do veículo",
            [
                f"Com o veículo {placa} já em uso às 07:30 de 30/07/2026, tentar outro agendamento "
                "no mesmo horário → conflito detectado e bloqueado.",
            ],
        ),
        (
            "6. Cartão do Motorista",
            [
                f"Sem agendamento válido para {n_a}, o botão/impressão do cartão não aparece (ou é bloqueado).",
                f"No cartão da viagem 30/07/2026 conferir: {n_mot}, {n_a}, {ac_nome}, "
                "Hospital Estadual / Campinas, 07:30, observações e IDADE (paciente e AC) coerentes "
                "com a data de nascimento na data da viagem.",
            ],
        ),
        (
            "7. Múltiplos cartões",
            [
                f"Programar {n_mot} com 5 viagens usando pacientes já cadastrados "
                f"({n_a}, {n_b}, {n_c}, {n_d} + 1 extra) e imprimir: 5 cartões, A4 paisagem, "
                "um abaixo do outro, corte fácil e espaçamento uniforme.",
            ],
        ),
        (
            "8. Layout do cartão",
            [
                f"Comparar o cartão de {n_a} com ModeloCartaoQuePreciso.jpeg "
                "(tamanho, margens, fontes, alinhamentos, dados, assinatura).",
            ],
        ),
        (
            "9. Folha Espelho",
            [
                f"Na Folha Espelho de 30/07/2026 conferir: {n_mot}, {placa}, {frota_txt}, "
                "rota Hospital Estadual, pacientes/acompanhantes, IDADE / IDADE AC, horários e total.",
            ],
        ),
        (
            "10. Ordem da impressão",
            [
                f"Ordem dos cartões: {n_a} → {n_b} → {n_c} → {n_d} "
                "(mesma ordem do agendamento); Folha Espelho na mesma sequência.",
            ],
        ),
        (
            "11. Impressão",
            [
                "Imprimir/visualizar cartões + Folha Espelho em Chrome, Edge e Firefox "
                "(A4 paisagem, sem quebra, sem corte de texto, sem páginas extras).",
            ],
        ),
        (
            "12. Informações obrigatórias",
            [
                "Teste negativo: programar viagem sem motorista (ou sem destino) e tentar imprimir — "
                "sistema destaca, informa e impede impressão se o dado for indispensável.",
            ],
        ),
        (
            "13. Casos extremos",
            [
                f"{n_b} (SEM acompanhante): cartão imprime normalmente, campo acompanhante vazio/omitido.",
                f"{n_a} + {ac_nome}: cartão traz o acompanhante; motorista com várias viagens gera vários cartões.",
            ],
        ),
        (
            "14. Consistência",
            [
                f"Alterar o horário de {n_a} de 07:30 para 08:00 → cartão e Folha Espelho atualizam; "
                f"excluir o agendamento de {n_d} → some do cartão e da folha.",
            ],
        ),
        (
            "15. UX",
            [
                "Na tela do agendamento/programação da viagem 30/07/2026: botões Imprimir, Visualizar e Voltar "
                "funcionam; mensagens e confirmações aparecem quando aplicável.",
                f"Na edição de {n_a}: botão Salvar (paciente) ≠ Salvar novo(s) acompanhante(s); "
                "não exibir erro pedindo AC quando já existe acompanhante ativo.",
            ],
        ),
        (
            "16. Fluxo completo (End-to-End)",
            [
                f"Usar cadastros JÁ EXISTENTES: frota {frota_txt} → veículo {placa} → motorista {n_mot} → "
                f"BUSCAR paciente {n_a} → cadastrar/usar acompanhante {ac_nome} vinculado a ele → "
                "agendar → validar → Cartão → Folha → impressão A4.",
            ],
        ),
        (
            "17. O que foi implantado para validar hoje",
            [
                "Idade: módulo único validadores/idade.py; formulários sem campo Idade; "
                "idade só em listagens e documentos (Cartão / Folha Espelho / impressão).",
                f"Ficha do paciente {n_a}: UX de acompanhantes (formulario novo recolhido se já houver AC; "
                "mensagem amigável se salvar formulário vazio).",
                f"Duplicata de {ac_nome} no paciente {n_a} desativada — conferir 1 ativo na ficha.",
                "Em Cadastro de Frota (aba Frota), o ℹ️ abre a orientação do fluxo frota → veículos.",
                "Ícone ℹ️ padronizado em Paciente, Acompanhante, Motorista e Veículos/Frota.",
                f"Após estas melhorias, reimprimir cartão de {n_a} e Folha Espelho de 30/07 — sem regressão.",
            ],
        ),
    ]


def build():
    global _BOX_ID
    _BOX_ID = 2000

    dados = carregar_dados_reais()
    secoes = montar_secoes(dados)

    doc = Document()
    set_narrow_margins(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    set_run_font(
        title.add_run("Checklist de Validação STP"),
        size=18,
        bold=True,
        color=(0x1A, 0x3A, 0x5C),
    )

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    set_run_font(
        sub.add_run("Impressão do Cartão do Motorista e Folha Espelho"),
        size=12,
        bold=True,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    set_run_font(
        meta.add_run(
            f"STP — Sistema de Transporte de Pacientes  |  Gerado em {datetime.now():%d/%m/%Y} "
            f"| Dados reais do banco local"
        ),
        size=9,
    )

    add_para(
        doc,
        "Objetivo: validar o fluxo desde o cadastro até a impressão do Cartão do Motorista e da "
        "Folha Espelho, incluindo regras de idade (só listagens/documentos) e UX de acompanhantes "
        "na ficha do paciente. Os exemplos usam registros JÁ CADASTRADOS no sistema. "
        "Para acompanhante: sempre BUSCAR o paciente existente e vincular — nunca inventar paciente.",
        size=10,
        space_after=6,
    )

    add_header_fields(doc)

    environment = doc.add_paragraph()
    environment.paragraph_format.space_after = Pt(8)
    add_clickable_checkbox(environment, "stp_ambiente_local")
    set_run_font(environment.add_run(" Local (localhost:5022)     "), size=10)
    add_clickable_checkbox(environment, "stp_ambiente_homolog")
    set_run_font(environment.add_run(" Homologação     "), size=10)
    add_clickable_checkbox(environment, "stp_ambiente_prod")
    set_run_font(environment.add_run(" Produção"), size=10)

    add_exemplo_cadastro(doc, dados)

    n = 1
    for section_title, items in secoes:
        heading = doc.add_heading(section_title, level=2)
        style_heading(heading, 2)
        n = add_checklist_table(doc, items, start_num=n)
        doc.add_paragraph()

    add_conclusao(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Gerado: {OUT}")
    print(f"Itens: {n - 1}")
    print(f"Paciente A: {dados['pacientes'][0]['nome'] if dados['pacientes'] else '—'}")
    print(f"Acompanhantes no banco: {dados['total_acompanhantes']}")
    print(f"Tamanho: {OUT.stat().st_size} bytes")
    return OUT


if __name__ == "__main__":
    build()
