# -*- coding: utf-8 -*-
"""Gera AJUSTES/Checklist/Checklist_Auditoria_Agendamentos_STP.docx

Checklist exclusivo e editável para auditoria completa do módulo Agendamentos
(coração do STP). Sem alterações funcionais — apenas o documento de validação.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AJUSTES" / "Checklist" / "Checklist_Auditoria_Agendamentos_STP.docx"

_BOX_ID = 5000


def _next_box_id() -> int:
    global _BOX_ID
    _BOX_ID += 1
    return _BOX_ID


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shading)


def set_run_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, *, size=11, bold=False, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_narrow_margins(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    return p


def add_clickable_checkbox(paragraph, name: str, box_id: int | None = None) -> None:
    if box_id is None:
        box_id = _next_box_id()
    safe_name = (
        name.replace("&", "&amp;")
        .replace('"', "&quot;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    xml = (
        '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
        "<w:sdtPr>"
        f'<w:alias w:val="{safe_name}"/>'
        f'<w:tag w:val="{safe_name}"/>'
        f'<w:id w:val="{box_id}"/>'
        "<w14:checkbox>"
        '<w14:checked w14:val="0"/>'
        '<w14:checkedState w14:val="2612" w14:font="Segoe UI Symbol"/>'
        '<w14:uncheckedState w14:val="2610" w14:font="Segoe UI Symbol"/>'
        "</w14:checkbox>"
        "</w:sdtPr>"
        "<w:sdtContent>"
        "<w:r><w:rPr>"
        '<w:rFonts w:ascii="Segoe UI Symbol" w:hAnsi="Segoe UI Symbol"/>'
        '<w:sz w:val="20"/>'
        f"</w:rPr><w:t>☐</w:t></w:r>"
        "</w:sdtContent>"
        "</w:sdt>"
    )
    paragraph._p.append(parse_xml(xml))


def add_checklist_table(doc, items, start_num=1, nota_padrao=""):
    """Tabela: # | Item | OK | Não OK | Revisar | Observações."""
    headers = ["#", "Item de validação", "OK", "Não OK", "Revisar", "Observações"]
    widths = [Cm(0.8), Cm(9.0), Cm(1.2), Cm(1.5), Cm(1.5), Cm(3.8)]
    table = doc.add_table(rows=1 + len(items), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), size=8, bold=True, color=(255, 255, 255))
        set_cell_shading(hdr[i], "1A3A5C")
        hdr[i].width = widths[i]

    checkbox_cols = {2: "ok", 3: "nao_ok", 4: "revisar"}

    for idx, text in enumerate(items):
        if isinstance(text, (tuple, list)):
            item_text, obs = text[0], (text[1] if len(text) > 1 else nota_padrao)
        else:
            item_text, obs = text, nota_padrao
        item_num = start_num + idx
        row = table.rows[idx + 1].cells
        for c in range(6):
            row[c].text = ""
            p = row[c].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if c == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(p.add_run(str(item_num)), size=8, bold=True)
            elif c == 1:
                set_run_font(p.add_run(item_text), size=8)
            elif c in checkbox_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_clickable_checkbox(p, f"ag_chk{item_num}_{checkbox_cols[c]}")
            else:
                set_run_font(p.add_run(obs or ""), size=7)
            row[c].width = widths[c]
        if idx % 2 == 1:
            for c in range(6):
                set_cell_shading(row[c], "F0F4F8")

    return start_num + len(items)


def add_info_table(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    w0 = widths[0] if widths else Cm(5.0)
    w1 = widths[1] if widths else Cm(12.8)
    for i, (lab, val) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        set_run_font(c0.paragraphs[0].add_run(lab), size=9, bold=True)
        set_cell_shading(c0, "E8EEF4")
        c0.width = w0
        c1.text = ""
        set_run_font(c1.paragraphs[0].add_run(val), size=9)
        c1.width = w1
    return table


def build():
    doc = Document()
    set_narrow_margins(doc.sections[0])
    n = 1

    add_para(
        doc,
        "STP — Sistema de Transporte de Pacientes",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_para(
        doc,
        "CHECKLIST EXCLUSIVO — Auditoria Profissional do Módulo de Agendamentos",
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_para(
        doc,
        "Coração da aplicação · Rota: /transporte/agendamentos",
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} · Documento editável · "
        "Nenhuma alteração funcional deve ser feita sem concluir a validação deste checklist.",
        size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    # ----- 1. Identificação -----
    add_heading(doc, "1. Identificação da auditoria", level=1)
    add_info_table(
        doc,
        [
            ("Responsável técnico:", "_________________________________"),
            ("Responsável operacional (SMS):", "_________________________________"),
            ("Data início:", "____/____/________"),
            ("Data fim:", "____/____/________"),
            ("Ambiente:", "☐ Local (127.0.0.1:5022)   ☐ Homologação   ☐ Produção"),
            ("WhatsApp nesta auditoria:", "☐ BLOQUEADO (obrigatório) — STP_BLOQUEAR_WHATSAPP=1"),
            ("Versão / commit:", "_________________________________"),
            ("Resultado geral:", "☐ Aprovado   ☐ Aprovado com ressalvas   ☐ Reprovado"),
        ],
    )

    add_heading(doc, "2. Princípios e escopo", level=1)
    add_para(
        doc,
        "A rota /transporte/agendamentos concentra o fluxo operacional da Secretaria Municipal "
        "de Saúde: cadastro → programação → impressão (Cartão / Folha Espelho) → status → "
        "(futuro) WhatsApp. Esta auditoria valida UI, regras de negócio, cenários, impressão, "
        "programação, lote, sandbox WhatsApp e logs — sem disparo real de mensagens.",
        size=9,
        space_after=6,
    )
    add_para(doc, "Pré-condições obrigatórias antes de iniciar:", size=10, bold=True, space_after=3)
    n = add_checklist_table(
        doc,
        [
            "Servidor STP reiniciado com código atual (use_reloader=False — reinício manual).",
            "STP_BLOQUEAR_WHATSAPP=1 ativo (iniciar_STP.ps1) — zero envio real ao telefone.",
            "Login com usuário de teste (ex.: admin) — anotar nome completo exibido na Folha Espelho.",
            "Modelo oficial Folha Espelho e Cartão do Motorista disponíveis para comparação visual.",
            "Seeds de teste conhecidos OU dados reais de homologação documentados (IDs anotados).",
            "Nenhuma alteração de código durante a execução dos testes (apenas registro de achados).",
        ],
        start_num=n,
    )

    # ----- 3. UI -----
    add_heading(doc, "3. Auditoria de Interface (UI) — /agendamentos", level=1)
    add_para(
        doc,
        "Avaliar na listagem, filtros, cards mobile, formulários Novo/Corrigir/Programar e páginas de impressão.",
        size=9,
        space_after=4,
    )
    n = add_checklist_table(
        doc,
        [
            "Alinhamentos de colunas da tabela (ID, Data/Hora, Paciente, Especialidade, Rota, Motorista, Veículo/Frota, Status, Ações).",
            "Espaçamentos internos/externos consistentes com o restante do STP.",
            "Responsividade: tabela desktop + cards mobile (stp-list-desktop / stp-list-mobile).",
            "Organização visual do cabeçalho (título, Folha Espelho, Excluir selecionados).",
            "Legibilidade de textos longos (truncamento + tooltip em paciente/rota/motorista).",
            "Contraste de status (Agendado, Confirmado, Em Andamento, Concluído, Cancelado).",
            "Consistência dos ícones de ação (editar, programar, confirmar, iniciar, concluir, cancelar, cartão).",
            "Estados dos botões: habilitado / desabilitado / hover (Folha Espelho e Cartão).",
            "Badges/indicadores “Aguardando” quando sem motorista/veículo/frota.",
            "Filtros: ID, busca, paciente, datas, status, motorista, destino, origem, placa, especialidade, pills de período.",
            "Paginação e per_page preservam filtros na URL.",
            "Ordenação padrão (data/hora desc) é compreensível para o atendente.",
            "Pesquisa geral (q) encontra paciente/destino/motorista/origem.",
            "Indicador visual de Folha Espelho liberada vs bloqueada (mensagem clara).",
            "Mensagens flash de sucesso/erro/aviso após criar, programar, mudar status, excluir.",
            "Estado vazio / nenhum resultado com filtros — mensagem amigável.",
            "Loading/feedback em ações demoradas (quando aplicável).",
            "Acessibilidade: title/aria-label nos ícones; checkbox “selecionar página”; contraste mínimo.",
            "Coluna ID sempre visível e legível (identificação operacional).",
        ],
        start_num=n,
    )

    # ----- 4. Funcionalidades -----
    add_heading(doc, "4. Auditoria de Funcionalidades", level=1)
    n = add_checklist_table(
        doc,
        [
            "Criação (Etapa 1 — /agendamentos/novo): paciente, especialidade, data/hora, origem, destino.",
            "Edição cadastral (Corrigir) antes da programação.",
            "Edição cadastral após programação (o que pode / o que não pode mudar).",
            "Exclusão individual (se existir) e exclusão em massa com confirmação.",
            "Cancelamento de status (ícone) e efeito na listagem (cancelados ocultos por padrão).",
            "Reativação de cancelado (existe? se não, registrar como ausência / decisão de negócio).",
            "Alteração de paciente no Corrigir.",
            "Alteração de acompanhante no Corrigir / formulário de agendamento.",
            "Alteração de motorista na Programar.",
            "Alteração de veículo na Programar.",
            "Alteração de frota na Programar (1 frota = 1 veículo).",
            "Alteração de especialidade.",
            "Alteração de origem e destino.",
            "Alteração de horários (saída / consulta se houver).",
            "Alterações após programação: Folha/Cartão refletem dados novos ao reimprimir.",
            "Alterações antes da programação: não liberam Folha/Cartão indevidamente.",
            "Fluxo de status: agendado → confirmado → em_andamento → concluido.",
            "Bloqueio de iniciar transporte sem programação (motorista + veículo/frota).",
        ],
        start_num=n,
    )

    # ----- 5. Regras de negócio -----
    add_heading(doc, "5. Regras de Negócio", level=1)
    n = add_checklist_table(
        doc,
        [
            "Campos obrigatórios no cadastro (paciente, especialidade, data, hora, origem, destino).",
            "Paciente “Necessita acompanhante” exige AC cadastrado antes de agendar.",
            "Programação exige motorista + (veículo OU frota).",
            "Folha Espelho liberada somente com programação completa (veículo ou frota).",
            "Cartão do Motorista liberado com a mesma regra de programação.",
            "Status cancelado impede impressão.",
            "Consistência status × ações disponíveis na linha.",
            "Permissões por perfil (atendente/supervisor/admin) — o que cada um pode fazer.",
            "Dependências entre campos (tipo recurso veículo vs frota).",
            "Atendente na Folha/Cartão = usuário logado no momento da impressão.",
            "Idade do paciente/AC calculada (não digitável) nos documentos.",
            "Consistência banco × interface (motorista_id, veiculo_id, frota_id exibidos corretamente).",
            "Conflito de horário do mesmo motorista (existe bloqueio? documentar resultado).",
            "Conflito de horário do mesmo veículo/frota (existe bloqueio? documentar resultado).",
            "Capacidade do veículo vs pacientes+acompanhantes (se implementado).",
        ],
        start_num=n,
        nota_padrao="",
    )

    # ----- 6. Cenários -----
    add_heading(doc, "6. Checklist exclusivo — Cenários de fluxo completo", level=1)
    add_para(
        doc,
        "Para cada cenário: anotar IDs gerados, data/hora, motorista, veículo/frota e resultado da impressão.",
        size=9,
        space_after=4,
    )

    cenarios = [
        (
            "6.1 Cenário 1 — Paciente simples",
            [
                "Cadastrar paciente sem acompanhante (ou usar existente sem AC).",
                "Criar 1 agendamento (1 viagem).",
                "Programar 1 motorista + 1 veículo.",
                "Confirmar status e imprimir Cartão (somente=1) + Folha Espelho.",
                "Validar campos: paciente, rota, horários, motorista, veículo, atendente logado.",
            ],
        ),
        (
            "6.2 Cenário 2 — Paciente com acompanhante",
            [
                "Paciente com AC ativo na ficha; selecionar AC no agendamento.",
                "Programar e imprimir Cartão — AC aparece (nome/parentesco/RG/tel conforme layout).",
                "Folha Espelho exibe AC e idades coerentes.",
                "Listagem mostra vínculo paciente/AC corretamente na edição.",
            ],
        ),
        (
            "6.3 Cenário 3 — Mesmo motorista, várias viagens no mesmo dia",
            [
                "Programar 3+ viagens no mesmo dia para o mesmo motorista.",
                "Abrir Cartões do dia — 1 cartão por viagem, ordem por horário.",
                "Layout A4 paisagem · até 3 por folha · corte/espaçamento.",
                "Verificar conflitos de horário sobrepostos (sistema bloqueia ou alerta?).",
            ],
        ),
        (
            "6.4 Cenário 4 — Mesmo veículo, várias viagens",
            [
                "Programar 2+ viagens no mesmo veículo com horários próximos/sobrepostos.",
                "Registrar se há bloqueio de conflito de veículo.",
                "Impressão Folha Espelho lista apenas elegíveis programados.",
            ],
        ),
        (
            "6.5 Cenário 5 — Motorista com destinos/cidades diferentes",
            [
                "Mesmo motorista: destinos em cidades distintas no mesmo dia.",
                "Ordem lógica na listagem e nos cartões do dia.",
                "Programação e impressão sem misturar dados entre viagens.",
            ],
        ),
        (
            "6.6 Cenário 6 — Diversos pacientes, mesmo destino",
            [
                "Vários pacientes para o mesmo destino/data.",
                "Filtro por destino/data funciona.",
                "Folha Espelho / cartões não agrupam indevidamente dados de pacientes.",
            ],
        ),
        (
            "6.7 Cenário 7 — Diversos pacientes, destinos diferentes",
            [
                "Mix de destinos na mesma página de listagem.",
                "Impressão filtrada por destino isolado.",
                "Sem vazamento de dados entre folhas.",
            ],
        ),
        (
            "6.8 Cenário 8 — Cancelamentos",
            [
                "Cancelar agendamento programado.",
                "Sumiu da listagem padrão (filtro status cancelado para ver).",
                "Cartão/Folha indisponíveis após cancelar.",
                "WhatsApp de cancelamento NÃO enviado (modo bloqueado) — log/simulação ok.",
            ],
        ),
        (
            "6.9 Cenário 9 — Reprogramações",
            [
                "Alterar Programar: trocar motorista mantendo veículo/frota.",
                "Alterar horário e reimprimir — documentos atualizados.",
                "Reprogramar após Em Andamento (permitido? documentar).",
            ],
        ),
        (
            "6.10 Cenário 10 — Alterações após já existir programação",
            [
                "Corrigir origem/destino com programação existente.",
                "Reimprimir Folha/Cartão reflete alteração.",
                "Status e ações permanecem coerentes.",
            ],
        ),
        (
            "6.11 Cenário 11 — Remover motorista",
            [
                "É possível limpar motorista na Programar? Se sim: Folha/Cartão devem bloquear.",
                "Se não for possível limpar: documentar regra e UX.",
            ],
        ),
        (
            "6.12 Cenário 12 — Trocar veículo",
            [
                "Trocar veículo na Programar; Folha/Cartão atualizam placa/frota.",
                "Trocar de veículo para frota (e vice-versa).",
            ],
        ),
        (
            "6.13 Cenário 13 — Trocar acompanhante",
            [
                "Trocar AC no Corrigir; Cartão/Folha mostram novo AC.",
                "Remover AC quando paciente não exige — documentos ok.",
            ],
        ),
        (
            "6.14 Cenário 14 — Trocar paciente",
            [
                "Trocar paciente no Corrigir; validações de AC/condição aplicadas.",
                "Documentos impressos refletem novo paciente.",
            ],
        ),
        (
            "6.15 Cenário 15 — Trocar especialidade",
            [
                "Alterar especialidade; aparece corretamente em listagem e documentos.",
            ],
        ),
        (
            "6.16 Cenário 16 — Trocar destino",
            [
                "Alterar destino; Folha Espelho e Cartão atualizam DESTINO.",
            ],
        ),
        (
            "6.17 Cenário 17 — Trocar horário",
            [
                "Alterar hora de saída; cartões do dia reordenam; Folha atualiza H. SAIDA.",
            ],
        ),
        (
            "6.18 Cenário 18 — Grande volume",
            [
                "Listagem com centenas de registros: tempo de carregamento aceitável.",
                "Filtros e paginação responsivos.",
                "Impressão “Todas” / intervalos 1–2 / 1–3 não travam o navegador.",
                "Exclusão em massa em página grande (cuidado — confirmar contagem).",
            ],
        ),
    ]

    for titulo, itens in cenarios:
        add_heading(doc, titulo, level=2)
        n = add_checklist_table(doc, itens, start_num=n)

    # ----- 7. Impressão -----
    add_heading(doc, "7. Auditoria da Impressão — Cartão do Motorista", level=1)
    n = add_checklist_table(
        doc,
        [
            "Um cartão (somente=1) — layout A4 paisagem.",
            "Vários cartões do mesmo motorista/data (lote do dia).",
            "Motorista com 1 viagem.",
            "Motorista com 2 viagens.",
            "Motorista com 3 viagens (1 folha).",
            "Motorista com muitas viagens (múltiplas folhas, quebra a cada 3).",
            "Quebra correta entre folhas; sem página em branco indevida.",
            "Orientação paisagem; margens; alinhamentos; fontes.",
            "Campos: motorista, frota/veículo, H. saída, destino, paciente, idade, AC, H. consulta, ponto, tel, rua, tratamento, atendente.",
            "Datas e horários corretos.",
            "Acompanhantes / pacientes / destinos / origem sem sobreposição ou corte.",
            "QR Code / códigos (se existirem no modelo — senão marcar N/A em Observações).",
            "Comparação visual com ModeloCartaoQuePreciso.jpeg / cartaomotorista.jpeg.",
            "Botão Imprimir na página gerada; Voltar funcional.",
            "Ícone Cartão na listagem: turquesa quando liberado; cinza quando bloqueado.",
        ],
        start_num=n,
    )

    add_heading(doc, "8. Auditoria da Impressão — Folha Espelho", level=1)
    n = add_checklist_table(
        doc,
        [
            "Layout idêntico ao modelo oficial (AJUSTES/folhaespelho.jpeg).",
            "Liberada na listagem quando há ≥1 agendamento com motorista + veículo/frota.",
            "Bloqueada quando só há “Aguardando”.",
            "Página atual / Pág. 1–2 / Pág. 1–3 / Todas — intervalos corretos.",
            "Rota individual /agendamentos/<id>/folha-espelho (tela Programar).",
            "Quebra automática / múltiplas páginas se volume alto.",
            "Todos os campos oficiais preenchidos (incl. ATENDENTE = usuário logado).",
            "Negritos, tamanhos de fonte, margens, espaçamentos.",
            "Consistência visual com o modelo da Secretaria.",
            "Não inclui cancelados nem não programados.",
            "Frota-only (sem veiculo_id) também libera Folha Espelho.",
        ],
        start_num=n,
    )

    # ----- 9. Programação -----
    add_heading(doc, "9. Auditoria da Programação", level=1)
    n = add_checklist_table(
        doc,
        [
            "Programação manual (tela Programar): veículo ou frota + motorista + observações.",
            "Após salvar, permanece na Programar com impressão liberada.",
            "Botões Folha/Cartão desabilitados antes de salvar programação.",
            "Alteração da programação existente.",
            "Remoção/limpeza da programação (se suportado) — senão documentar.",
            "Conflitos / sobreposição de horários (motorista e veículo).",
            "Disponibilidade do motorista.",
            "Disponibilidade do veículo/frota.",
            "Programação automática (se existir no sistema — senão N/A).",
        ],
        start_num=n,
    )

    # ----- 10. Massa -----
    add_heading(doc, "10. Auditoria das Impressões em Massa", level=1)
    n = add_checklist_table(
        doc,
        [
            "Impressão individual (Cartão ícone / Folha do ID / Programar).",
            "Impressão em lote Folha Espelho (intervalos de páginas).",
            "Impressão filtrada (data, status, paciente).",
            "Impressão por motorista (Cartões do dia).",
            "Impressão por data (filtro + Folha).",
            "Impressão por cidade/destino (filtro destino).",
            "Impressão por veículo/placa (filtro placa).",
            "Seleção + imprimir marcados (existe? se não, registrar gap / prioridade).",
        ],
        start_num=n,
    )

    # ----- 11. WhatsApp sandbox -----
    add_heading(doc, "11. Integração WhatsApp — modo teste (sandbox)", level=1)
    add_para(
        doc,
        "PROIBIDO envio real nesta auditoria. Validar apenas infraestrutura com STP_BLOQUEAR_WHATSAPP=1.",
        size=9,
        bold=True,
        space_after=4,
    )
    n = add_checklist_table(
        doc,
        [
            "Flag STP_BLOQUEAR_WHATSAPP=1 impede iniciar serviço/lembretes no boot.",
            "agendar_mensagem / _enviar_mensagem_agora registram [SIMULAÇÃO] e não abrem WhatsApp Web.",
            "Geração de mensagem de confirmação (template + variáveis) pode ser inspecionada sem envio.",
            "Placeholders: nome, data, hora, origem, destino, especialidade.",
            "Formatação da mensagem legível.",
            "Destinatário resolvido (tel_cel preferencial) — sem disparo.",
            "Anexos (N/A se não houver).",
            "Logs em logs_whatsapp/ (sucesso/erro) — comportamento sob bloqueio documentado.",
            "Fila de envio não processa mensagens reais com bloqueio ativo.",
            "Tratamento de telefone inválido.",
            "Após 100% de aprovação futura: plano para remover bloqueio só em produção.",
            "Criar agendamento novo NÃO envia zap real com bloqueio ativo.",
            "Cancelar NÃO envia zap real com bloqueio ativo.",
        ],
        start_num=n,
    )

    # ----- 12. Logs -----
    add_heading(doc, "12. Auditoria de Logs / Trilha", level=1)
    n = add_checklist_table(
        doc,
        [
            "Criação de agendamento registrada (console/log/auditoria).",
            "Edição cadastral: snapshot/auditoria de alteração (se implementado).",
            "Exclusão / exclusão em massa.",
            "Programação salva.",
            "Impressão (Cartão/Folha) — existe log de emissão? (hoje: tipicamente NÃO — registrar gap).",
            "Cancelamento / reprogramação / mudança de status.",
            "Tentativa de WhatsApp sob sandbox ([SIMULAÇÃO] no console).",
            "Erros e exceções com mensagem útil ao atendente + log técnico.",
        ],
        start_num=n,
    )

    # ----- 13. Pré-achados conhecidos -----
    add_heading(doc, "13. Pré-achados técnicos (referência — validar no teste)", level=1)
    add_para(
        doc,
        "Itens já observados em auditorias anteriores. Confirmar no ambiente reiniciado e atualizar status.",
        size=9,
        space_after=4,
    )
    n = add_checklist_table(
        doc,
        [
            (
                "Folha Espelho na listagem libera com frota (sem veiculo_id) — correção recente.",
                "Validar c/ ID programado só em frota",
            ),
            (
                "ATENDENTE na Folha/Cartão = usuário logado (não texto fixo JULIANA).",
                "Validar imprimindo logado",
            ),
            (
                "Impressão na tela Programar (liberada só após salvar programação).",
                "Validar fluxo",
            ),
            (
                "Coluna ID + filtro ID na listagem.",
                "Validar UI",
            ),
            (
                "Comprovante de agendamento (balcão) — AUSENTE.",
                "Gap — priorizar após esta auditoria",
            ),
            (
                "Log de emissão de documentos — AUSENTE.",
                "Gap",
            ),
            (
                "Export CSV/Excel da listagem — AUSENTE.",
                "Gap",
            ),
            (
                "Seleção + imprimir Folha/Cartões marcados — AUSENTE (checkbox só exclui).",
                "Gap",
            ),
            (
                "Conflito motorista/veículo no mesmo horário — confirmar se bloqueia de fato.",
                "Validar Cenários 3–4",
            ),
            (
                "WhatsApp sandbox via env — ativo no iniciar_STP.ps1.",
                "Manter durante auditoria",
            ),
        ],
        start_num=n,
    )

    # ----- 14. Relatório final -----
    add_heading(doc, "14. Relatório final obrigatório (preencher após os testes)", level=1)
    add_para(doc, "14.1 Contagem", size=10, bold=True, space_after=3)
    add_info_table(
        doc,
        [
            ("Itens aprovados (OK):", "________"),
            ("Itens reprovados (Não OK):", "________"),
            ("Itens para revisar:", "________"),
            ("Decisão final:", "☐ Aprovado  ☐ Aprovado c/ ressalvas  ☐ Reprovado"),
        ],
    )

    add_para(doc, "14.2 Inconsistências e riscos", size=10, bold=True, space_before=8, space_after=3)
    add_para(
        doc,
        "Listar cada inconsistência com: descrição · evidência (ID/print) · risco · impacto operacional / "
        "atendentes / motoristas / Secretaria · prioridade (Crítica/Alta/Média/Baixa) · correção sugerida.",
        size=8,
        space_after=4,
    )

    headers = ["#", "Inconsistência / risco", "Prioridade", "Impacto", "Correção sugerida", "Status"]
    widths = [Cm(0.7), Cm(5.5), Cm(2.0), Cm(3.5), Cm(4.0), Cm(2.1)]
    table = doc.add_table(rows=1 + 12, cols=6)
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), size=8, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, "1A3A5C")
        cell.width = widths[i]
    for r in range(1, 13):
        for c in range(6):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            if c == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(p.add_run(str(r)), size=8, bold=True)
            else:
                set_run_font(p.add_run(""), size=8)
            cell.width = widths[c]
        if r % 2 == 0:
            for c in range(6):
                set_cell_shading(table.rows[r].cells[c], "F0F4F8")

    add_para(doc, "14.3 Melhorias recomendadas (após aprovação dos testes)", size=10, bold=True, space_before=10, space_after=3)
    for _ in range(6):
        add_para(doc, "• ___________________________________________________________________________", size=9, space_after=4)

    add_para(doc, "14.4 Assinaturas", size=10, bold=True, space_before=8, space_after=4)
    add_info_table(
        doc,
        [
            ("Auditor técnico:", "_________________________  Data: ____/____/________"),
            ("Responsável SMS / operação:", "_________________________  Data: ____/____/________"),
            ("Autorização para correções:", "☐ Sim, iniciar correções priorizadas   ☐ Aguardar"),
        ],
    )

    add_para(
        doc,
        "Objetivo final: módulo de Agendamentos validado, robusto e preparado para produção, "
        "concentrando o fluxo operacional do STP com segurança (WhatsApp só após 100% dos testes).",
        size=8,
        space_before=10,
        space_after=4,
    )
    add_para(
        doc,
        "Arquivo: AJUSTES/Checklist/Checklist_Auditoria_Agendamentos_STP.docx · "
        "Script gerador: scripts/gerar_checklist_auditoria_agendamentos_stp.py",
        size=7,
        space_after=2,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")
    print(f"Total de itens de checklist (aprox. numeração final): {n - 1}")
    return OUT


if __name__ == "__main__":
    build()
