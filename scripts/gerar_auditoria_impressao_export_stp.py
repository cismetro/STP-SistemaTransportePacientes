# -*- coding: utf-8 -*-
"""Gera AJUSTES/Auditoria_Impressao_Visualizacao_Exportacao_STP.docx

Auditoria de impressão, visualização, PDF, exportação e relatórios do STP.
Documento editável para revisão e devolução.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AJUSTES" / "Checklist" / "Auditoria_Impressao_Visualizacao_Exportacao_STP.docx"


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shading)


def set_run_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, *, size=11, bold=False, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_narrow_margins(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    return p


def fill_header_row(cells, headers, widths, fill="1A3A5C"):
    for i, h in enumerate(headers):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), size=8, bold=True, color=(255, 255, 255))
        set_cell_shading(cells[i], fill)
        if widths:
            cells[i].width = widths[i]


def add_table(doc, headers, rows, widths=None, zebra=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    fill_header_row(table.rows[0].cells, headers, widths)
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            set_run_font(p.add_run(str(val)), size=8)
            if widths:
                cells[c_idx].width = widths[c_idx]
        if zebra and r_idx % 2 == 1:
            for c in cells:
                set_cell_shading(c, "F0F4F8")
    return table


def add_review_table(doc, items):
    """Tabela editável: # | Item | Prioridade | Decisão | Observações do revisor."""
    headers = ["#", "Item", "Prioridade", "Decisão (manter/alterar/adiar)", "Observações do revisor"]
    widths = [Cm(0.8), Cm(6.5), Cm(2.0), Cm(4.0), Cm(4.5)]
    table = doc.add_table(rows=1 + len(items), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    fill_header_row(table.rows[0].cells, headers, widths)
    for idx, (item, pri) in enumerate(items):
        cells = table.rows[idx + 1].cells
        vals = [str(idx + 1), item, pri, "", ""]
        for c, v in enumerate(vals):
            cells[c].text = ""
            p = cells[c].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            set_run_font(p.add_run(v), size=8, bold=(c == 0))
            cells[c].width = widths[c]
        if idx % 2 == 1:
            for c in cells:
                set_cell_shading(c, "F0F4F8")
    return table


def build():
    doc = Document()
    set_narrow_margins(doc.sections[0])

    add_para(
        doc,
        "STP — Sistema de Transporte de Pacientes",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_para(
        doc,
        "AUDITORIA COMPLETA — Impressão, Visualização, PDF, Exportação e Relatórios",
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} · Escopo: monólito app.py (ativo) · "
        "Sem implementação nesta etapa — documento para edição e devolução",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    # Campos de revisão
    add_heading(doc, "1. Dados da revisão", level=1)
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    campos = [
        ("Responsável pela revisão:", "_________________________________"),
        ("Data da revisão:", "____/____/________"),
        ("Versão do sistema / commit:", "_________________________________"),
        ("Decisão geral:", "☐ Aprovar onda 1   ☐ Aprovar parcial   ☐ Revisar prioridades"),
    ]
    for i, (lab, val) in enumerate(campos):
        c0, c1 = meta.rows[i].cells
        c0.text = ""
        set_run_font(c0.paragraphs[0].add_run(lab), size=9, bold=True)
        set_cell_shading(c0, "E8EEF4")
        c0.width = Cm(5.0)
        c1.text = ""
        set_run_font(c1.paragraphs[0].add_run(val), size=9)
        c1.width = Cm(12.8)

    add_heading(doc, "2. Veredito executivo", level=1)
    add_para(
        doc,
        "O fluxo operacional diário depende de documentos físicos e digitais. No monólito ativo "
        "(app.py), a impressão é feita via HTML + window.print() (sem auto-print). Os documentos "
        "oficiais Folha Espelho e Cartão do Motorista estão maduros. Não há PDF/CSV nas listagens, "
        "comprovante de agendamento para o balcão, log de emissão nem impressão na tela Programar. "
        "O export Excel de backup é stub (capa vazia). O pacote legado sistema/ possui export "
        "PDF/Excel/CSV, mas não está ligado ao create_app() ativo.",
        size=10,
        space_after=8,
    )
    add_para(doc, "Resumo quantitativo (auditoria):", size=10, bold=True, space_after=4)
    add_table(
        doc,
        ["Indicador", "Valor"],
        [
            ["Módulos / telas com impressão OK", "Pacientes, Acompanhantes, Motoristas, Veículos, Cartão, Folha"],
            ["Sem impressão relevante", "Frotas, Programar, Dashboard, Faturamento, Uso, Comprovante"],
            ["Itens críticos", "2"],
            ["Itens altos", "4"],
            ["Itens médios", "5"],
            ["Itens baixos", "3"],
        ],
        widths=[Cm(7.0), Cm(10.8)],
    )

    add_heading(doc, "3. Itens já implementados", level=1)
    add_table(
        doc,
        ["#", "Recurso", "Onde / como"],
        [
            ["1", "Shell de impressão A4 paisagem", "montar_shell_impressao + gerar_folhas_tabela"],
            ["2", "Botões padrão de listagem (4 ranges)", "gerar_botoes_impressao — pacientes, acompanhantes, motoristas, veículos"],
            ["3", "Impressão listagem Pacientes", "/pacientes/imprimir — HTML tabular com filtros"],
            ["4", "Impressão listagem Acompanhantes", "/acompanhantes/imprimir — idades; sem data nascimento"],
            ["5", "Impressão listagem Motoristas", "/motoristas/imprimir"],
            ["6", "Impressão listagem Veículos", "/veiculos/imprimir (aba veículo)"],
            ["7", "Folha Espelho (documento oficial)", "/agendamentos/imprimir — só programados elegíveis"],
            ["8", "Cartão do Motorista (lote do dia)", "/agendamentos/<id>/cartao-motorista — 3/folha A4"],
            ["9", "Pré-visualização em nova aba", "target=_blank + botão Imprimir na página gerada"],
            ["10", "Disabled + tooltip (Cartão / Folha)", "agendamento_elegivel_* + gerar_botoes_folha_espelho"],
            ["11", "Relatórios gerenciais (abas)", "/relatorios — print da tela (window.print)"],
            ["12", "Print de tela combustível", "/combustivel/relatorio"],
        ],
        widths=[Cm(0.8), Cm(6.0), Cm(11.0)],
    )

    add_heading(doc, "4. Cobertura por módulo", level=1)
    add_para(
        doc,
        "Legenda status: OK = funciona; Parcial = existe com lacuna; Ausente = não existe; Stub = existe mas vazio/inútil.",
        size=9,
        space_after=6,
    )
    add_table(
        doc,
        ["Módulo", "Impressão", "Export", "PDF", "Ficha/doc", "Detalhe", "Sua nota"],
        [
            ["Pacientes (listagem)", "OK", "Ausente", "Ausente", "Ausente", "HTML A4 · pág. atual/1–2/1–3/todas", ""],
            ["Acompanhantes (listagem)", "OK", "Ausente", "Ausente", "Ausente", "Mesmo padrão das listagens", ""],
            ["Motoristas (listagem)", "OK", "Ausente", "Ausente", "Ausente", "HTML A4 tabular", ""],
            ["Veículos (aba veículo)", "OK", "Ausente", "Ausente", "Ausente", "HTML A4 tabular", ""],
            ["Frotas (aba frota)", "Ausente", "Ausente", "Ausente", "Ausente", "Listagem sem botões de impressão", ""],
            ["Agendamentos — Folha Espelho", "Parcial", "Ausente", "Ausente", "OK", "Enable UI ignora frota-only", ""],
            ["Agendamentos — Cartão Motorista", "OK", "Ausente", "Ausente", "OK", "Lote motorista+data · 3/folha", ""],
            ["Agendamentos — Programar", "Ausente", "Ausente", "Ausente", "Ausente", "Sem Imprimir/Visualizar após programar", ""],
            ["Comprovante de agendamento", "Ausente", "Ausente", "Ausente", "Ausente", "Não existe documento de balcão", ""],
            ["Dashboard", "Ausente", "Ausente", "Ausente", "Ausente", "Cards/agenda sem print/export", ""],
            ["Relatórios (/relatorios)", "Parcial", "Ausente", "Ausente", "Parcial", "Só print da aba ativa; período fraco", ""],
            ["Faturamento", "Ausente", "Ausente", "Ausente", "Ausente", "Detalhes sem imprimir fatura", ""],
            ["Uso de veículos", "Ausente", "Ausente", "Ausente", "Ausente", "Detalhes sem print", ""],
            ["Combustível / relatório", "Parcial", "Ausente", "Ausente", "Parcial", "Print de tela; NF é campo texto", ""],
            ["Backup Excel", "Ausente", "Stub", "Ausente", "Ausente", "openpyxl gera capa sem dados", ""],
        ],
        widths=[Cm(4.2), Cm(1.6), Cm(1.6), Cm(1.4), Cm(1.6), Cm(4.8), Cm(2.6)],
    )

    add_heading(doc, "5. Agendamentos — matriz de documentos", level=1)
    add_table(
        doc,
        ["Documento / recurso", "Status", "Observação", "Sua nota"],
        [
            ["Folha Espelho", "OK", "Oficial · HTML · nova aba · só programados", ""],
            ["Cartão do Motorista", "OK", "Lote motorista+data · 3/folha", ""],
            ["Comprovante de agendamento", "Ausente", "Não implementado", ""],
            ["Relação diária", "Ausente", "Indireto via filtros da listagem", ""],
            ["Relação por motorista", "Parcial", "Cartão cobre lote do dia; sem relatório tabular", ""],
            ["Relação por veículo / frota", "Ausente", "Filtro placa existe; sem documento próprio", ""],
            ["Relação por especialidade / destino / data", "Parcial", "Filtros na listagem; sem layout de relação", ""],
            ["Impressão na tela Programar", "Ausente", "Checklist espera; código não tem", ""],
            ["Histórico de emissão", "Ausente", "Sem tabela/log de documentos emitidos", ""],
            ["Reimpressão com via / trilha", "Ausente", "Hoje = reabrir a mesma URL", ""],
        ],
        widths=[Cm(5.5), Cm(1.8), Cm(7.5), Cm(3.0)],
    )

    add_para(doc, "Bug confirmado (Folha Espelho × frota):", size=10, bold=True, space_before=8, space_after=4)
    add_para(
        doc,
        "Na listagem, tem_programados exige veiculo_id IS NOT NULL e motorista_id, ignorando frota_id. "
        "Já agendamento_tem_recurso_programado / elegibilidade do Cartão aceitam frota. "
        "Programação só com frota pode deixar os botões da Folha Espelho desabilitados na UI, "
        "mesmo com o documento teoricamente elegível.",
        size=9,
        space_after=8,
    )

    add_heading(doc, "6. Cadastros individuais (ficha)", level=1)
    add_table(
        doc,
        ["Cadastro", "Impressão do registro único?", "Necessidade sugerida", "Sua nota"],
        [
            ["Paciente (editar)", "Não", "Ficha completa / resumida", ""],
            ["Acompanhante", "Não", "Ficha vinculada ao paciente", ""],
            ["Motorista", "Não", "Ficha + status CNH", ""],
            ["Veículo", "Não", "Ficha + situação / frota", ""],
            ["Frota", "Não", "Cadastro + veículo vinculado", ""],
            ["Agendamento (novo/corrigir)", "Não", "Comprovante / resumo", ""],
            ["Usuário (visualizar)", "Não", "Baixa prioridade", ""],
        ],
        widths=[Cm(4.5), Cm(3.5), Cm(5.5), Cm(4.3)],
    )

    add_heading(doc, "7. Histórico, reimpressão e compartilhamento", level=1)
    add_table(
        doc,
        ["Tema", "Situação atual", "Proposta", "Sua nota"],
        [
            [
                "Histórico de emissão",
                "Ausente (sem quem/quando/tipo)",
                "Log: documento, agendamento, usuário, data/hora, via",
                "",
            ],
            [
                "Reimpressão",
                "Indireta (reabrir URL)",
                "Botão Reimprimir + contador de vias",
                "",
            ],
            [
                "Compartilhamento",
                "Ausente (WhatsApp = lembrete, não documento)",
                "Avaliar PDF/link interno; não obrigatório na onda 1",
                "",
            ],
        ],
        widths=[Cm(3.5), Cm(5.0), Cm(5.5), Cm(3.8)],
    )

    add_heading(doc, "8. Padronização UX dos botões", level=1)
    add_table(
        doc,
        ["Padrão", "Onde", "Formato", "Sua nota"],
        [
            ["A — Listagens", "Pacientes, AC, Motoristas, Veículos", "Texto “Imprimir:” + 4 botões print-btn", ""],
            ["B — Folha Espelho", "Agendamentos", "“Folha Espelho:” + 4 botões OU chip disabled", ""],
            ["C — Cartão", "Linha do agendamento", "Ícone ti-id · tooltip · disabled sem href", ""],
            ["D — Relatórios", "/relatorios, combustível", "Botão Imprimir na própria página (emoji)", ""],
            ["E — Documento oficial", "Cartão / Folha gerados", "Toolbar Imprimir + Voltar", ""],
        ],
        widths=[Cm(3.5), Cm(5.0), Cm(6.5), Cm(2.8)],
    )
    add_para(
        doc,
        "Recomendação: unificar em um componente (rótulo + ranges + disabled) e ícone Tabler "
        "consistente; Visualizar = mesma URL sem auto-print.",
        size=9,
        space_after=8,
    )

    add_heading(doc, "9. Priorização — ausentes, incompletos e melhorias", level=1)
    add_para(
        doc,
        "Use a coluna “Decisão” e “Observações do revisor” para editar e devolver.",
        size=9,
        space_after=6,
    )

    add_heading(doc, "9.1 Crítico", level=2)
    add_review_table(
        doc,
        [
            (
                "Comprovante de agendamento (paciente/balcão) — atendente precisa entregar no cadastro; "
                "hoje só Folha/Cartão pós-programação",
                "CRÍTICO",
            ),
            (
                "Imprimir / Visualizar na tela Programar após salvar — fluxo diário programa → imprime na hora; "
                "hoje redireciona à listagem",
                "CRÍTICO",
            ),
        ],
    )

    add_heading(doc, "9.2 Alto", level=2)
    add_review_table(
        doc,
        [
            (
                "Corrigir enable Folha Espelho alinhado à frota (tem_programados deve aceitar frota_id)",
                "ALTO",
            ),
            (
                "Export CSV/Excel das listagens e de agendamentos do dia",
                "ALTO",
            ),
            (
                "Relação diária (por data / motorista / veículo) como documento",
                "ALTO",
            ),
            (
                "Seleção + imprimir Folha/Cartões dos marcados (checkbox hoje só exclui)",
                "ALTO",
            ),
        ],
    )

    add_heading(doc, "9.3 Médio", level=2)
    add_review_table(
        doc,
        [
            ("Impressão da aba Frotas (paridade com veículos)", "MÉDIO"),
            ("Ficha individual (paciente / motorista / veículo)", "MÉDIO"),
            ("Impressão de fatura e detalhe de uso de veículo", "MÉDIO"),
            ("PDF real + export em /relatorios (período filtrando todas as abas)", "MÉDIO"),
            ("Log de emissão + reimpressão com via", "MÉDIO"),
        ],
    )

    add_heading(doc, "9.4 Baixo / opcional", level=2)
    add_review_table(
        doc,
        [
            ("Padronizar botões (ícone + tooltip + Visualizar explícito)", "BAIXO"),
            ("Completar Excel de backup / ocultar demo cartão em produção", "BAIXO"),
            ("Dashboard: resumo diário imprimível (passagem de plantão)", "BAIXO"),
        ],
    )

    add_heading(doc, "10. Ondas sugeridas (após aprovação)", level=1)
    add_para(doc, "Onda 1 — Crítico / Alto (operação diária)", size=10, bold=True, space_after=2)
    add_para(
        doc,
        "Comprovante de agendamento · Imprimir/Visualizar na Programar · Corrigir enable Folha com frota · "
        "Relação diária · Export CSV agendamentos/listagens · Imprimir selecionados",
        size=9,
        space_after=6,
    )
    add_para(doc, "Onda 2 — Médio (paridade e arquivo)", size=10, bold=True, space_after=2)
    add_para(
        doc,
        "Impressão frotas · Fichas individuais · Fatura/uso · PDF em relatórios · Log de emissão",
        size=9,
        space_after=6,
    )
    add_para(doc, "Onda 3 — Baixo (polimento)", size=10, bold=True, space_after=2)
    add_para(
        doc,
        "Unificar UX dos botões · Completar Excel backup · Ocultar demo cartão · Resumo dashboard imprimível",
        size=9,
        space_after=10,
    )

    add_heading(doc, "11. Espaço livre para anotações do revisor", level=1)
    for _ in range(8):
        add_para(doc, "_" * 95, size=10, space_after=8)

    add_para(
        doc,
        "Fonte técnica: app.py (gerar_html_impressao_*, Folha Espelho, Cartão Motorista, /relatorios, "
        "backup Excel stub) · Legado sistema/ não ativo no create_app · Nenhuma funcionalidade foi "
        "alterada nesta auditoria.",
        size=8,
        space_before=6,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")
    return OUT


if __name__ == "__main__":
    build()
