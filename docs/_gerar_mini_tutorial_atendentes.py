# -*- coding: utf-8 -*-
"""Mini tutorial STP — Atendentes (pacientes, acompanhantes, motoristas, veículos, agendamentos)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(
    r"D:\Projetos\python\STP-SistemaTransportePacientes\docs"
    r"\Mini_Tutorial_STP_Atendentes_Cadastros_Agendamentos.docx"
)


def set_run(run, *, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x6E, 0x6E)


def para(doc, text, *, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold, size=size)
    p.paragraph_format.space_after = Pt(6)


def bullet(doc, text):
    bp = doc.add_paragraph(text, style="List Bullet")
    for run in bp.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, ht in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ht
        shade(cell, "0D6E6E")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
                run.font.name = "Calibri"
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
    doc.add_paragraph()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("STP — Sistema de Transporte de Pacientes")
    set_run(r, bold=True, size=18, color=RGBColor(0x0D, 0x6E, 0x6E))

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(
        "Mini tutorial para Atendentes\n"
        "Cadastro de pacientes, acompanhantes, motoristas, veículos e agendamentos\n"
        "Prefeitura Municipal de Cosmópolis"
    )
    set_run(r, bold=True, size=12)

    para(
        doc,
        "Este guia ensina o dia a dia do Atendente: cadastrar o que é necessário e "
        "agendar o transporte em duas etapas (pedido → programação com veículo/motorista).",
    )

    h(doc, "1. Para quem é este tutorial", 1)
    para(doc, "Usuários com perfil Atendente (operações básicas):")
    table(
        doc,
        ["ID", "Nome", "Usuário", "Perfil", "Situação"],
        [
            ["8", "Alisson Fernandes", "alisson_fernandes", "Atendente", "Ativo"],
            ["7", "Vitor Reinaldo Endrighi Aleixo", "vitor_aleixo", "Atendente", "Ativo"],
            ["6", "Jane Maria Baccarin", "jane_maria", "Atendente", "Ativo"],
            ["5", "Collins Alexandre Pereira", "collins_pereira", "Atendente", "Ativo"],
            ["4", "Stefany Barreto Cunha", "stefany_barreto", "Atendente", "Ativo"],
        ],
    )
    para(doc, "O Atendente pode usar:", bold=True)
    bullet(doc, "Cadastros: Pacientes, Acompanhantes, Motoristas, Veículos")
    bullet(doc, "Agendamentos (criar, programar, acompanhar status, impressões)")
    bullet(doc, "Início (dashboard), Frota (consulta de uso/combustível) e Relatórios")
    para(doc, "O Atendente NÃO acessa:", bold=True)
    bullet(doc, "Usuários, WhatsApp e Faturamento (áreas de administrador / supervisor / contador)")

    h(doc, "2. Como entrar no sistema", 1)
    table(
        doc,
        ["Item", "Valor"],
        [
            ["Endereço", "http://127.0.0.1:5022/transporte"],
            ["Tela de login", "http://127.0.0.1:5022/transporte/login"],
            ["Campos", "Usuário * e Senha * → botão Entrar"],
        ],
    )
    bullet(doc, "Use o usuário liberado para você (ex.: jane_maria, alisson_fernandes).")
    bullet(doc, "Se não lembrar a senha, peça ao administrador para redefinir.")

    h(doc, "3. Menu principal (o que você vai usar)", 1)
    table(
        doc,
        ["Menu", "Para quê"],
        [
            ["Agendamentos", "Criar pedido e programar viagem (tela principal do dia a dia)"],
            ["Início", "Visão geral / dashboard"],
            ["Cadastros → Pacientes", "Cadastrar e editar pacientes"],
            ["Cadastros → Acompanhantes", "Vincular acompanhante a um paciente"],
            ["Cadastros → Motoristas", "Cadastrar motoristas da frota"],
            ["Cadastros → Veículos", "Cadastrar veículo e/ou frota"],
            ["Relatórios", "Consultas e impressões gerenciais"],
        ],
    )

    h(doc, "4. Ordem correta do trabalho (importante!)", 1)
    para(
        doc,
        "Siga esta ordem. Evita erro na hora de agendar:",
        bold=True,
    )
    table(
        doc,
        ["Passo", "O que fazer", "Por quê"],
        [
            ["1", "Cadastrar o Paciente", "Sem paciente não há agendamento nem acompanhante"],
            ["2", "Cadastrar Acompanhante (se precisar)", "Obrigatório quando a condição exige acompanhante"],
            ["3", "Ter Motorista e Veículo/Frota cadastrados", "Usados na 2ª etapa (Programar Transporte)"],
            ["4", "Novo Agendamento (etapa 1)", "Registra o pedido: paciente, data, origem e destino"],
            ["5", "Programar Transporte (etapa 2)", "Escolhe veículo/frota + motorista"],
            ["6", "Confirmar / Iniciar / Concluir", "Acompanha o status da viagem"],
        ],
    )

    h(doc, "5. Cadastrar Paciente", 1)
    para(doc, "Caminho: Cadastros → Pacientes → Cadastrar Novo Paciente")
    table(
        doc,
        ["Campo", "Obrigatório?", "Dica"],
        [
            ["Nome Completo", "Sim", "Nome completo, sem abreviação excessiva"],
            ["CPF", "Sim", "CPF válido e único no sistema"],
            ["Telefone Celular / Residencial", "Pelo menos um", "Prefira celular com DDD"],
            ["Data de Nascimento", "Sim", "Confira o ano"],
            ["Condição do paciente", "Não", "Se marcar, escolha a condição; se Outros, descreva"],
            ["CEP", "Sim", "Use a busca ViaCEP"],
            ["Logradouro / Número / Bairro", "Sim", "Endereço de residência"],
            ["Ponto de Embarque", "Sim", "Onde o carro busca o paciente"],
            ["Cartão SUS", "Não", "Se preencher, deve ser CNS válido"],
            ["Necessidades Especiais / Observações", "Não", "Alergias, cadeira de rodas, etc."],
        ],
    )
    bullet(doc, "Botão: Salvar Paciente.")
    bullet(
        doc,
        "Se a condição for “Necessita acompanhante”, o sistema sugere cadastrar o acompanhante em seguida — faça isso antes de agendar.",
    )

    h(doc, "6. Cadastrar Acompanhante", 1)
    para(doc, "Caminho: Cadastros → Acompanhantes → Cadastrar Acompanhante")
    para(doc, "Regra do sistema: “Sem paciente não há acompanhante.”", bold=True)
    table(
        doc,
        ["Campo", "Obrigatório?", "Dica"],
        [
            ["Paciente", "Sim", "Escolha o paciente já cadastrado e ativo"],
            ["Parentesco", "Não", "Se Outros, especifique"],
            ["Nome", "Sim*", "Ou marque “Nome ainda não informado / desconhecido”"],
            ["RG", "Não", "Se informar, o sistema valida"],
            ["Telefone", "Não", "Útil para contato no dia"],
            ["Data de Nascimento", "Sim", "—"],
        ],
    )
    bullet(doc, "Pode usar “Cadastrar mais um” para vários acompanhantes do mesmo paciente.")
    bullet(doc, "Botão: Salvar acompanhante(s).")

    h(doc, "7. Cadastrar Motorista", 1)
    para(doc, "Caminho: Cadastros → Motoristas → Cadastrar Novo Motorista")
    table(
        doc,
        ["Campo", "Obrigatório?", "Dica"],
        [
            ["Nome Completo", "Sim", "—"],
            ["CPF", "Sim", "Único no sistema"],
            ["Telefone", "Sim", "Com DDD"],
            ["Data de Nascimento", "Sim", "—"],
            ["Número da CNH", "Sim", "Única no sistema"],
            ["Categoria CNH", "Sim", "A–E, AB–AE etc."],
            ["Vencimento da CNH", "Sim", "Não deixe CNH vencida sem atualizar"],
            ["CEP / Logradouro / Número / Bairro / Endereço", "Sim", "Use CEP para preencher"],
            ["Status", "Sim", "Ativo / Inativo / Férias / Licença"],
            ["Observações", "Não", "—"],
        ],
    )
    bullet(doc, "Para programar viagem, o motorista precisa estar disponível (em geral Status = Ativo).")

    h(doc, "8. Cadastrar Veículo e Frota", 1)
    para(doc, "Caminho: Cadastros → Veículos")
    para(doc, "Há duas abas: Veículo e Frota. Regra: 1 frota = 1 veículo.", bold=True)

    h(doc, "8.1 Aba Veículo", 2)
    table(
        doc,
        ["Campo", "Obrigatório?", "Dica"],
        [
            ["Placa", "Sim", "Placa do veículo"],
            ["Tipo de Veículo", "Sim", "Ambulância, Van, Micro-ônibus, Carro"],
            ["Marca / Modelo / Ano", "Sim", "Pode usar consulta FIPE para ajudar"],
            ["Cor", "No formulário", "Paleta de cores"],
            ["Capacidade de Passageiros", "Sim", "Pode vir da FIPE"],
            ["Adaptado para PCD", "Não", "Sim / Não"],
            ["Observações", "Não", "—"],
        ],
    )

    h(doc, "8.2 Aba Frota", 2)
    table(
        doc,
        ["Campo", "Obrigatório?", "Exemplo"],
        [
            ["Número da Frota", "Sim", "F00267"],
            ["Nome da Frota", "Sim", "NI Frota 267"],
        ],
    )
    bullet(doc, "Depois de salvar a frota, vincule ou cadastre o veículo correspondente.")

    h(doc, "9. Agendamento — Etapa 1 (pedido)", 1)
    para(doc, "Caminho: Agendamentos → Novo Agendamento")
    para(
        doc,
        "Nesta etapa NÃO se escolhe veículo nem motorista. Só o pedido de transporte.",
        bold=True,
    )
    table(
        doc,
        ["Campo", "Obrigatório?", "Dica"],
        [
            ["Paciente", "Sim", "Filtrar por ID, nome ou CPF → Filtrar"],
            ["Especialidade", "Sim", "Se Outro, informe o texto"],
            ["Levará acompanhante nesta viagem", "Se marcado / se condição exige", "Escolha o acompanhante cadastrado"],
            ["Data", "Sim", "Data da viagem"],
            ["Hora de Saída", "Sim", "Horário de saída do veículo"],
            ["Hora da Consulta", "Não", "Aparece na Folha Espelho / Cartão"],
            ["Origem", "Sim*", "Pode marcar busca em endereço diferente do cadastro"],
            ["Destino", "Sim", "CEP / Cidade de SP / Destino Predefinido / Manual"],
        ],
    )
    para(doc, "Destino — escolha um modo:", bold=True)
    bullet(doc, "Buscar por CEP → CEP, cidade e endereço/local de destino")
    bullet(doc, "Cidade de SP → cidade de São Paulo + endereço na cidade")
    bullet(doc, "Destino Predefinido → cidade + local CNES ou texto livre")
    bullet(doc, "Manual → endereço completo")
    bullet(doc, "Botão: Cadastrar Agendamento.")
    bullet(doc, "Na listagem o item pode aparecer como “○ Sem programação” até a etapa 2.")

    h(doc, "10. Agendamento — Etapa 2 (programar)", 1)
    para(doc, "Caminho: Agendamentos → na linha do pedido → Programar Transporte (ou Alterar Programação)")
    table(
        doc,
        ["Campo", "Obrigatório?", "Dica"],
        [
            ["Recurso da viagem", "Sim", "Veículo OU Frota"],
            ["Veículo ou Frota", "Sim", "Conforme a opção escolhida"],
            ["Motorista", "Sim", "Motorista ativo"],
            ["Observações", "Não", "Ex.: levar cadeira de rodas"],
        ],
    )
    bullet(doc, "Botão: Salvar Programação.")
    bullet(doc, "Depois disso o agendamento fica “● Programado”.")

    h(doc, "11. Status da viagem (depois de programar)", 1)
    table(
        doc,
        ["Status", "Significado / ação"],
        [
            ["Agendado", "Pedido criado (padrão)"],
            ["Confirmado", "Confirmado com o paciente / operação"],
            ["Em Andamento", "Viagem iniciada (só se já estiver programado)"],
            ["Concluído", "Viagem finalizada (status travado)"],
            ["Cancelado", "Cancelado — use Reativar para voltar a Agendado, se necessário"],
        ],
    )
    bullet(doc, "A listagem padrão esconde cancelados; filtre Status = Cancelado para vê-los.")
    bullet(doc, "Impressões úteis: Folha Espelho e Cartão do Motorista (indisponíveis se cancelado).")

    h(doc, "12. Exemplo rápido do dia a dia", 1)
    para(doc, "Situação: paciente precisa ir a consulta em Campinas amanhã.")
    bullet(doc, "1) Confirme se o paciente existe em Cadastros → Pacientes (senão, cadastre).")
    bullet(doc, "2) Se precisa de acompanhante, cadastre/selecione em Acompanhantes.")
    bullet(doc, "3) Agendamentos → Novo Agendamento → preencha paciente, especialidade, data, hora, origem e destino → Cadastrar.")
    bullet(doc, "4) Na lista, Programar Transporte → escolha frota/veículo + motorista → Salvar.")
    bullet(doc, "5) No dia: Confirmar → Iniciar → Concluir; imprima Folha Espelho/Cartão se a operação pedir.")

    h(doc, "13. Problemas comuns", 1)
    table(
        doc,
        ["Problema", "O que fazer"],
        [
            ["Não consigo salvar paciente", "Confira CPF válido, telefone e endereço completo"],
            ["Sistema pede acompanhante", "Cadastre o acompanhante no paciente antes de agendar"],
            ["Não acho o botão de veículo no Novo Agendamento", "É normal: veículo/motorista só na etapa Programar"],
            ["Não consigo colocar Em Andamento", "Programe veículo/frota + motorista primeiro"],
            ["Não vejo agendamento cancelado", "Filtre Status = Cancelado na listagem"],
            ["Esqueci a senha", "Peça ao administrador (menu Usuários)"],
        ],
    )

    h(doc, "14. Checklist do Atendente", 1)
    for item in [
        "☐ Entrei com meu usuário de Atendente",
        "☐ Paciente cadastrado (CPF, telefone, embarque)",
        "☐ Acompanhante cadastrado quando necessário",
        "☐ Motorista e veículo/frota disponíveis (ou já existiam)",
        "☐ Etapa 1: Novo Agendamento salvo",
        "☐ Etapa 2: Programar Transporte salvo",
        "☐ Status atualizado e impressões feitas se preciso",
    ]:
        bp = doc.add_paragraph(item)
        for run in bp.runs:
            run.font.size = Pt(11)

    rod = doc.add_paragraph()
    rod.paragraph_format.space_before = Pt(16)
    r = rod.add_run(
        "Documento interno STP — Cosmópolis. Destinado aos Atendentes listados neste tutorial. "
        "Em caso de dúvida de permissão ou senha, contate o administrador do sistema."
    )
    set_run(r, size=9, color=RGBColor(0x66, 0x66, 0x66))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Gerado: {OUT}")


if __name__ == "__main__":
    main()
